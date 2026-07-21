#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
prajna 企智 - 招投标助手生成器 v1.0.0

为销售/投标团队一键生成招投标套件：
1. Excel 工作簿（6 个工作表）
   - 招标文件拆解清单
   - 投标人资格自审表
   - 技术偏离表
   - 商务报价表
   - 评分响应索引
   - 时间节点与材料清单
2. Word 投标文件大纲

依赖：openpyxl、python-docx
"""

import argparse
import os
import re
import sys
from datetime import datetime, timedelta
from pathlib import Path

try:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter
except ImportError as exc:  # pragma: no cover
    print("[错误] 缺少 openpyxl，请执行：pip install openpyxl", file=sys.stderr)
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError as exc:  # pragma: no cover
    print("[错误] 缺少 python-docx，请执行：pip install python-docx", file=sys.stderr)
    sys.exit(1)


SKILL_NAME = "prajna-bidding-assistant"
DEFAULT_OUTPUT_DIR = Path.home() / ".prajna" / SKILL_NAME / "samples"

THEME_BLUE = "1F4E78"
THEME_LIGHT_BLUE = "D9E1F2"
THEME_BG = "F2F2F2"

HEADER_FILL = PatternFill(start_color=THEME_BLUE, end_color=THEME_BLUE, fill_type="solid")
HEADER_FONT = Font(name="微软雅黑", size=11, bold=True, color="FFFFFF")
SUBHEADER_FILL = PatternFill(start_color=THEME_LIGHT_BLUE, end_color=THEME_LIGHT_BLUE, fill_type="solid")
SECTION_FONT = Font(name="微软雅黑", size=12, bold=True, color=THEME_BLUE)
NORMAL_FONT = Font(name="微软雅黑", size=10)
BOLD_FONT = Font(name="微软雅黑", size=10, bold=True)
BLUE_FONT = Font(name="微软雅黑", size=10, color="0000FF")
GRAY_FONT = Font(name="微软雅黑", size=10, color="808080")
RED_FONT = Font(name="微软雅黑", size=10, color="FF0000")
GREEN_FONT = Font(name="微软雅黑", size=10, color="008000")
ORANGE_FONT = Font(name="微软雅黑", size=10, color="FF6600")
CENTER_ALIGN = Alignment(horizontal="center", vertical="center", wrap_text=True)
LEFT_ALIGN = Alignment(horizontal="left", vertical="center", wrap_text=True)
THIN_BORDER = Border(
    left=Side(style="thin", color="B7B7B7"),
    right=Side(style="thin", color="B7B7B7"),
    top=Side(style="thin", color="B7B7B7"),
    bottom=Side(style="thin", color="B7B7B7"),
)

DISCLAIMER = "【人工智能生成-需人工核验】"


def _safe_filename(name):
    name = re.sub(r"[\\/:*?\"<>|\s]+", "_", str(name))
    return name.strip("_") or "project"


def _today():
    return datetime.now().strftime("%Y%m%d")


def _parse_amount(amount):
    try:
        return float(amount) if amount else 0.0
    except (TypeError, ValueError):
        return 0.0


def _industry_defaults(industry):
    industry = str(industry or "通用").strip()
    defaults = {
        "IT": {
            "items": [
                ("智能硬件与传感器", "套", 120, 8000),
                ("网络与安防设备", "套", 1, 980000),
                ("平台软件与中间件", "套", 1, 1200000),
                ("系统集成与实施", "项", 1, 1500000),
                ("三年运维服务", "年", 3, 360000),
                ("培训与知识转移", "人·天", 120, 2500),
            ],
            "tax_rate": 0.06,
        },
        "建筑": {
            "items": [
                ("土建工程", "项", 1, 2200000),
                ("装饰装修工程", "项", 1, 980000),
                ("安装工程", "项", 1, 1500000),
                ("设备及主材", "批", 1, 860000),
                ("措施项目费", "项", 1, 260000),
            ],
            "tax_rate": 0.09,
        },
        "服务": {
            "items": [
                ("年度人力资源外包服务", "年", 1, 1800000),
                ("招聘流程外包 RPO", "人", 80, 8000),
                ("薪酬社保代缴服务", "人·月", 240, 350),
                ("管理咨询与培训", "项", 1, 560000),
                ("驻场项目经理", "人·月", 12, 22000),
            ],
            "tax_rate": 0.06,
        },
        "制造": {
            "items": [
                ("生产线设备 A", "台", 3, 580000),
                ("生产线设备 B", "台", 2, 760000),
                ("辅助设备及工具", "批", 1, 420000),
                ("安装调试", "项", 1, 380000),
                ("操作培训", "人·天", 60, 2000),
                ("质保期内备件", "批", 1, 180000),
            ],
            "tax_rate": 0.13,
        },
    }
    return defaults.get(industry, defaults["IT"])


def _split_amount(amount, industry):
    """根据行业和金额，按比例拆分报价项，确保合计与输入金额一致。"""
    cfg = _industry_defaults(industry)
    base_items = cfg["items"]
    tax_rate = cfg["tax_rate"]

    if amount <= 0:
        amount = sum(q * p for _, _, q, p in base_items)

    base_total = sum(q * p for _, _, q, p in base_items)
    ratio = amount / base_total

    items = []
    for name, unit, qty, price in base_items:
        unit_price = round(price * ratio, 2)
        total = round(unit_price * qty, 2)
        items.append([name, unit, qty, unit_price, total, tax_rate])

    # 修正因逐项四舍五入导致的合计差额，将差额加在最后一项
    current_total = sum(item[4] for item in items)
    diff = round(amount - current_total, 2)
    if items and diff != 0:
        items[-1][4] = round(items[-1][4] + diff, 2)
        if items[-1][2] != 0:
            items[-1][3] = round(items[-1][4] / items[-1][2], 2)

    return [tuple(item) for item in items], tax_rate


def _project_scope(industry):
    scopes = {
        "IT": "系统平台建设、数据对接、设备供货、集成实施及运维服务。",
        "建筑": "工程施工、材料采购、安装调试及竣工验收。",
        "服务": "服务方案设计、人员派驻、流程执行、质量管控及持续优化。",
        "制造": "设备供货、安装调试、人员培训、试运行及质保服务。",
    }
    return scopes.get(industry, "按招标文件要求提供全部产品及服务。")


# ---------------------------------------------------------------------------
# Excel helpers
# ---------------------------------------------------------------------------
def _set_header(ws, row, values, start_col=1):
    for c, v in enumerate(values, start_col):
        cell = ws.cell(row=row, column=c, value=v)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = CENTER_ALIGN
        cell.border = THIN_BORDER


def _write_row(ws, row, values, is_input=False, is_note=False, is_warn=False, align=None):
    for c, v in enumerate(values, 1):
        cell = ws.cell(row=row, column=c, value=v)
        if is_input:
            cell.font = BLUE_FONT
        elif is_warn:
            cell.font = ORANGE_FONT
        elif is_note:
            cell.font = GRAY_FONT
        else:
            cell.font = NORMAL_FONT
        cell.alignment = align or (LEFT_ALIGN if isinstance(v, str) else CENTER_ALIGN)
        cell.border = THIN_BORDER


def _autofit_columns(ws, min_width=10, max_width=40):
    for column in ws.columns:
        max_length = 0
        col_letter = get_column_letter(column[0].column)
        for cell in column:
            if cell.value is not None:
                length = len(str(cell.value))
                if length > max_length:
                    max_length = length
        adjusted_width = min(max(max_length + 2, min_width), max_width)
        ws.column_dimensions[col_letter].width = adjusted_width


def _add_disclaimer(ws, row, text="本表由 prajna 企智自动生成，仅供辅助参考。"):
    cell = ws.cell(row=row, column=1, value=f"{DISCLAIMER} {text}")
    cell.font = GRAY_FONT
    cell.alignment = LEFT_ALIGN
    end_col = ws.max_column or 6
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=end_col)


# ---------------------------------------------------------------------------
# Excel sheet builders
# ---------------------------------------------------------------------------
def build_cover(wb, params):
    ws = wb.active
    ws.title = "项目信息"

    ws["B2"] = "招投标套件"
    ws["B2"].font = Font(name="微软雅黑", size=18, bold=True, color=THEME_BLUE)
    ws.merge_cells("B2:E2")

    ws["B4"] = "项目信息"
    ws["B4"].font = SECTION_FONT
    ws.merge_cells("B4:E4")

    info = [
        ("项目名称", params["project"]),
        ("投标人", params["bidder"] or "待填写"),
        ("招标人", params["tenderer"] or "待填写"),
        ("投标金额（元）", params["amount"]),
        ("工期/服务期", params["duration"] or "待填写"),
        ("行业领域", params["industry"] or "通用"),
        ("生成日期", datetime.now().strftime("%Y-%m-%d")),
    ]
    for i, (label, value) in enumerate(info, 5):
        ws.cell(row=i, column=2, value=label).font = BOLD_FONT
        ws.cell(row=i, column=2).alignment = LEFT_ALIGN
        ws.cell(row=i, column=3, value=value).font = BLUE_FONT
        ws.cell(row=i, column=3).alignment = LEFT_ALIGN

    ws["B13"] = "使用说明"
    ws["B13"].font = SECTION_FONT
    notes = [
        "1. 本套件由 prajna 企智 prajna-bidding-assistant 技能自动生成。",
        "2. 蓝色字体单元格为建议填写项，请根据实际招标文件修订。",
        "3. 所有数据仅供投标准备参考，正式投标前须人工复核。",
        "4. 报价表已自动拆分示例分项，实际报价请以招标文件为准。",
    ]
    for i, note in enumerate(notes, 14):
        ws.cell(row=i, column=2, value=note).font = GRAY_FONT
        ws.merge_cells(start_row=i, start_column=2, end_row=i, end_column=5)
        ws.cell(row=i, column=2).alignment = LEFT_ALIGN

    _add_disclaimer(ws, 19, "本套件内容不构成法律或商务承诺，正式投标须依据招标文件及企业实际情况调整。")
    _set_column_widths(ws, [4, 18, 40, 18, 18])
    return ws


def _set_column_widths(ws, widths):
    for col_idx, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(col_idx)].width = width


def build_tender_breakdown(wb, params):
    ws = wb.create_sheet("招标文件拆解清单")

    ws["A1"] = "一、招标文件拆解清单"
    ws["A1"].font = SECTION_FONT
    ws.merge_cells("A1:F1")

    headers = ["序号", "招标要求", "责任部门", "截止时间", "重要性", "备注"]
    _set_header(ws, 2, headers)
    ws.freeze_panes = "A3"

    rows = [
        ("投标截止时间与地点", "投标部", "T-0", "高", "务必提前 1 小时送达"),
        ("投标人资格要求", "法务/行政", "T-5", "高", "营业执照、资质证书、财务审计报告"),
        ("技术与服务要求", "技术/交付", "T-4", "高", "逐条响应，避免负偏离"),
        ("商务与报价要求", "商务/财务", "T-3", "高", "报价有效期、保证金、付款条件"),
        ("评分标准与权重", "投标负责人", "T-6", "高", "价格分/技术分/商务分拆解"),
        ("业绩与案例要求", "市场部", "T-5", "中", "近 3 年同类项目合同"),
        ("项目团队要求", "HR/项目", "T-4", "中", "项目经理及关键人员简历、证书"),
        ("质保与售后服务", "技术/客服", "T-3", "中", "质保期、响应时间、服务网点"),
        ("投标保证金", "财务", "T-2", "高", "金额、形式、退还条件"),
        ("合同条款与风险提示", "法务", "T-3", "中", "违约责任、付款节点、争议解决"),
    ]
    for i, row in enumerate(rows, 3):
        _write_row(ws, i, [i - 2] + list(row))

    _add_disclaimer(ws, len(rows) + 4)
    _autofit_columns(ws, min_width=10, max_width=36)
    ws.column_dimensions["B"].width = 36


def build_qualification(wb, params):
    ws = wb.create_sheet("投标人资格自审表")

    ws["A1"] = "二、投标人资格自审表"
    ws["A1"].font = SECTION_FONT
    ws.merge_cells("A1:F1")

    headers = ["序号", "资质项", "招标要求", "我方情况", "满足状态", "证明文件"]
    _set_header(ws, 2, headers)
    ws.freeze_panes = "A3"

    rows = [
        ("营业执照", "合法有效，经营范围覆盖", params["bidder"] + "（已注册）", "满足", "营业执照副本"),
        ("注册资本", "不少于招标要求", "1000 万元", "满足", "工商登记信息"),
        ("资质证书", "行业相关资质", "具备相关资质", "待核实", "资质证书扫描件"),
        ("财务审计报告", "近三年无亏损", "近三年审计报告", "待核实", "审计报告"),
        ("纳税证明", "近半年完税证明", "近半年纳税记录", "待核实", "完税证明"),
        ("社保缴纳", "近半年社保证明", "近半年社保记录", "待核实", "社保证明"),
        ("信用记录", "无重大违法失信记录", "信用中国查询截图", "待核实", "信用报告"),
        ("业绩案例", "近 3 年同类项目", "3 项以上", "待补充", "合同/验收报告"),
        ("项目经理", "具备相应资格证书", "1 名", "待补充", "简历/证书"),
        ("投标保证金", "按要求缴纳", "待缴纳", "待处理", "银行回单"),
    ]
    for i, row in enumerate(rows, 3):
        status = row[4]
        is_warn = status in ("待核实", "待补充", "待处理")
        _write_row(ws, i, [i - 2] + list(row), is_warn=is_warn)

    _add_disclaimer(ws, len(rows) + 4)
    _autofit_columns(ws, min_width=10, max_width=32)
    ws.column_dimensions["C"].width = 28
    ws.column_dimensions["D"].width = 28


def build_tech_deviation(wb, params):
    ws = wb.create_sheet("技术偏离表")

    ws["A1"] = "三、技术偏离表"
    ws["A1"].font = SECTION_FONT
    ws.merge_cells("A1:E1")

    headers = ["序号", "招标要求", "我方响应", "偏离情况", "说明"]
    _set_header(ws, 2, headers)
    ws.freeze_panes = "A3"

    scope = _project_scope(params["industry"])
    rows = [
        (f"项目总体要求：{scope}", f"完全响应，提供{params['industry'] or '相关'}整体解决方案", "无偏离", "详见技术方案第 2 章"),
        ("工期/服务期要求", params["duration"] or "按招标文件执行", "无偏离", "详见项目进度计划"),
        ("质量标准与验收要求", "满足国家及行业现行标准", "无偏离", "详见质量保证体系"),
        ("技术架构与平台要求", "支持主流技术栈，具备扩展能力", "无偏离", "详见技术架构设计"),
        ("数据安全与保密要求", "符合等级保护及数据安全法规", "无偏离", "详见安全方案"),
        ("接口与集成要求", "提供标准 API 及数据接口", "正偏离", "额外提供 SDK 与示例代码"),
        ("售后服务与响应时间", "7×24 小时响应，4 小时到场", "无偏离", "详见售后服务承诺"),
        ("培训与知识转移", "提供操作及维护培训", "正偏离", "额外提供视频教程及文档"),
        ("人员驻场要求", "按招标文件派驻人员", "待确认", "需根据最终需求确认"),
    ]
    for i, row in enumerate(rows, 3):
        deviation = row[3]
        is_warn = deviation in ("负偏离", "待确认")
        _write_row(ws, i, [i - 2] + list(row), is_warn=is_warn)

    _add_disclaimer(ws, len(rows) + 4)
    _autofit_columns(ws, min_width=10, max_width=38)
    ws.column_dimensions["B"].width = 38
    ws.column_dimensions["C"].width = 38


def build_commercial_quote(wb, params):
    ws = wb.create_sheet("商务报价表")

    ws["A1"] = "四、商务报价表"
    ws["A1"].font = SECTION_FONT
    ws.merge_cells("A1:G1")

    headers = ["序号", "分项名称", "单位", "数量", "单价（元）", "总价（元）", "税率"]
    _set_header(ws, 2, headers)
    ws.freeze_panes = "A3"

    items, default_tax = _split_amount(params["amount"], params["industry"])
    subtotal = 0
    for i, (name, unit, qty, price, total, tax) in enumerate(items, 3):
        _write_row(ws, i, [i - 2, name, unit, qty, price, total, f"{tax:.0%}"])
        ws.cell(row=i, column=5).number_format = "#,##0.00"
        ws.cell(row=i, column=6).number_format = "#,##0.00"
        subtotal += total

    total_row = len(items) + 3
    ws.cell(row=total_row, column=1, value="合计").font = BOLD_FONT
    ws.merge_cells(start_row=total_row, start_column=1, end_row=total_row, end_column=5)
    ws.cell(row=total_row, column=1).alignment = CENTER_ALIGN
    ws.cell(row=total_row, column=6, value=subtotal).font = BOLD_FONT
    ws.cell(row=total_row, column=6).number_format = "#,##0.00"
    ws.cell(row=total_row, column=7, value="—").font = BOLD_FONT
    for c in range(1, 8):
        ws.cell(row=total_row, column=c).border = THIN_BORDER

    ws.cell(row=total_row + 2, column=1, value=f"投标总价（人民币）：{subtotal:,.2f} 元（大写：{__rmb_capital(subtotal)}）")
    ws.cell(row=total_row + 2, column=1).font = BOLD_FONT
    ws.merge_cells(start_row=total_row + 2, start_column=1, end_row=total_row + 2, end_column=7)

    _add_disclaimer(ws, total_row + 4, "报价为示例拆分，正式投标须按招标文件报价清单及企业实际成本核算。")
    _autofit_columns(ws, min_width=10, max_width=32)
    ws.column_dimensions["B"].width = 32


def __rmb_capital(amount):
    """将数字金额转换为大写人民币（支持到亿级）。"""
    nums = "零壹贰叁肆伍陆柒捌玖"
    units = ["", "拾", "佰", "仟"]
    groups = ["", "万", "亿"]

    decimal = round(float(amount), 2)
    integer = int(decimal)
    fraction = round((decimal - integer) * 100)

    def _four_to_cap(n):
        """将 0-9999 的数字转为大写。"""
        if n == 0:
            return ""
        s = ""
        zero_flag = False
        for i in range(4):
            digit = n % 10
            n //= 10
            if digit == 0:
                if s and not zero_flag:
                    s = nums[0] + s
                    zero_flag = True
            else:
                s = nums[digit] + units[i] + s
                zero_flag = False
        # 去掉末尾的零（如 壹仟零壹拾 末尾无零，壹仟 末尾也无零）
        return s.rstrip(nums[0])

    def _int_to_cap(n):
        if n == 0:
            return nums[0]
        parts = []
        group_idx = 0
        while n > 0:
            part = n % 10000
            n //= 10000
            cap = _four_to_cap(part)
            if cap:
                if group_idx > 0:
                    # 处理类似 100000001 这种跨级补零
                    if parts and not parts[-1].endswith(nums[0]) and part < 1000:
                        cap = nums[0] + cap
                parts.append(cap + groups[group_idx])
            group_idx += 1
        return "".join(reversed(parts))

    int_part = _int_to_cap(integer)
    jiao = fraction // 10
    fen = fraction % 10
    if jiao > 0 or fen > 0:
        frac_part = ""
        if jiao > 0:
            frac_part += nums[jiao] + "角"
        if fen > 0:
            frac_part += nums[fen] + "分"
    else:
        frac_part = "整"

    return f"{int_part}元{frac_part}"


def build_scoring_index(wb, params):
    ws = wb.create_sheet("评分响应索引")

    ws["A1"] = "五、评分响应索引"
    ws["A1"].font = SECTION_FONT
    ws.merge_cells("A1:E1")

    headers = ["序号", "评分项", "分值", "我方响应要点", "证明材料位置"]
    _set_header(ws, 2, headers)
    ws.freeze_panes = "A3"

    rows = [
        ("价格部分", 30, "报价总价、分项报价、优惠条件", "商务报价表"),
        ("商务部分", 20, "资质、业绩、财务状况、信用记录", "资格自审表、资质文件"),
        ("技术方案", 25, "技术架构、实施路线、创新点", "技术偏离表、技术方案"),
        ("项目团队", 10, "项目经理经验、关键人员配置", "人员简历、资格证书"),
        ("服务承诺", 10, "售后响应、质保期限、本地化服务", "售后服务承诺函"),
        ("演示与答辩", 5, "现场演示、答疑能力", "演示 PPT、答辩提纲"),
    ]
    for i, row in enumerate(rows, 3):
        _write_row(ws, i, [i - 2] + list(row))

    ws.cell(row=len(rows) + 4, column=1, value="评分策略建议").font = SECTION_FONT
    ws.merge_cells(start_row=len(rows) + 4, start_column=1, end_row=len(rows) + 4, end_column=5)

    notes = [
        "1. 价格分占比最高，需在成本测算基础上制定有竞争力的报价策略。",
        "2. 技术方案应避免负偏离，对关键条款提供正偏离说明以争取加分。",
        "3. 业绩案例优先选择同类型、同区域、近 3 年内的项目。",
        "4. 所有证明材料须与评分项一一对应，便于评标专家快速定位。",
    ]
    for i, note in enumerate(notes, len(rows) + 5):
        ws.cell(row=i, column=1, value=note).font = GRAY_FONT
        ws.merge_cells(start_row=i, start_column=1, end_row=i, end_column=5)

    _add_disclaimer(ws, len(rows) + 10)
    _autofit_columns(ws, min_width=10, max_width=38)
    ws.column_dimensions["B"].width = 22
    ws.column_dimensions["D"].width = 38


def build_timeline(wb, params):
    ws = wb.create_sheet("时间节点与材料清单")

    ws["A1"] = "六、时间节点与材料清单"
    ws["A1"].font = SECTION_FONT
    ws.merge_cells("A1:F1")

    headers = ["阶段", "事项", "责任人", "截止日期", "交付物", "状态"]
    _set_header(ws, 2, headers)
    ws.freeze_panes = "A3"

    base = datetime.now()
    milestones = [
        ("招前研究", "招标文件研读与评分拆解", "投标负责人", base + timedelta(days=1), "招标文件拆解清单", "待启动"),
        ("资格自审", "资质材料收集与自审", "法务/行政", base + timedelta(days=2), "资格自审表", "待启动"),
        ("技术方案", "技术方案撰写与偏离分析", "技术/交付", base + timedelta(days=5), "技术偏离表", "待启动"),
        ("商务报价", "成本测算与报价编制", "商务/财务", base + timedelta(days=6), "商务报价表", "待启动"),
        ("评分响应", "评分项对应材料整理", "投标负责人", base + timedelta(days=7), "评分响应索引", "待启动"),
        ("内部评审", "法务、财务、技术联合评审", "投标负责人", base + timedelta(days=8), "评审记录", "待启动"),
        ("签章封装", "签字盖章、打印、装订、密封", "投标专员", base + timedelta(days=9), "投标文件", "待启动"),
        ("开标投标", "按时送达开标地点", "投标专员", base + timedelta(days=10), "送达回执", "待启动"),
    ]
    for i, (stage, item, owner, deadline, deliverable, status) in enumerate(milestones, 3):
        _write_row(ws, i, [stage, item, owner, deadline.strftime("%Y-%m-%d"), deliverable, status], is_warn=True)

    _add_disclaimer(ws, len(milestones) + 4)
    _autofit_columns(ws, min_width=10, max_width=32)
    ws.column_dimensions["B"].width = 32
    ws.column_dimensions["E"].width = 26


# ---------------------------------------------------------------------------
# Word document builder
# ---------------------------------------------------------------------------
def _add_heading(doc, text, level=1):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.font.name = "微软雅黑"
    run.font.size = Pt(16 if level == 1 else 14 if level == 2 else 12)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading


def _add_para(doc, text, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run.font.size = Pt(10.5)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    p.paragraph_format.line_spacing = 1.15
    return p


def build_word_outline(params, output_path):
    doc = Document()

    # Cover
    title = doc.add_paragraph()
    run = title.add_run("投标文件")
    run.font.name = "微软雅黑"
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    run = subtitle.add_run(f"项目名称：{params['project']}")
    run.font.name = "微软雅黑"
    run.font.size = Pt(14)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    cover_info = [
        ("投标人", params["bidder"] or "（待填写）"),
        ("招标人", params["tenderer"] or "（待填写）"),
        ("投标金额", f"{params['amount']:,.2f} 元" if params["amount"] else "（待填写）"),
        ("工期/服务期", params["duration"] or "（待填写）"),
        ("编制日期", datetime.now().strftime("%Y-%m-%d")),
    ]
    for label, value in cover_info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label}：")
        r1.font.name = "微软雅黑"
        r1.font.size = Pt(12)
        r2 = p.add_run(value)
        r2.font.name = "微软雅黑"
        r2.font.size = Pt(12)
        r2.bold = True

    doc.add_page_break()

    # TOC
    _add_heading(doc, "目录", level=1)
    toc_items = [
        "一、投标函",
        "二、法定代表人授权书",
        "三、投标人基本情况",
        "四、公司介绍与资质证明",
        "五、类似项目业绩",
        "六、技术方案",
        "七、商务方案",
        "八、报价文件",
        "九、售后服务承诺",
        "十、附件清单",
    ]
    for item in toc_items:
        _add_para(doc, item)

    doc.add_page_break()

    # Sections
    _add_heading(doc, "一、投标函", level=1)
    _add_para(doc, f"致：{params['tenderer'] or '招标人'}")
    _add_para(doc, f"我方 {params['bidder'] or '投标人'} 郑重声明：愿按招标文件要求参与 {params['project']} 的投标，并承诺所提交文件真实、有效。")
    _add_para(doc, "（此处为投标函正文模板，请根据招标文件要求填写完整内容。）")

    _add_heading(doc, "二、法定代表人授权书", level=1)
    _add_para(doc, "本授权书声明：注册于（公司地址）的（投标人名称）法定代表人（姓名）授权（被授权人姓名）为本公司的全权代表，代表本公司签署本次投标相关文件。")
    _add_para(doc, "（请附法定代表人及被授权人身份证明文件。）")

    _add_heading(doc, "三、投标人基本情况", level=1)
    _add_para(doc, "请填写投标人基本信息表，包括公司名称、注册地址、注册资本、成立时间、经营范围、联系方式、统一社会信用代码等。")

    _add_heading(doc, "四、公司介绍与资质证明", level=1)
    _add_heading(doc, "4.1 公司简介", level=2)
    _add_para(doc, "介绍公司发展历程、主营业务、核心能力、组织架构、服务网络等。")
    _add_heading(doc, "4.2 资质证书", level=2)
    _add_para(doc, "列示营业执照、行业资质、体系认证、荣誉证书等，并附扫描件。")

    _add_heading(doc, "五、类似项目业绩", level=1)
    _add_para(doc, "提供近 3 年同类项目业绩清单，包括项目名称、客户名称、合同金额、项目内容、验收情况、联系方式等。")

    _add_heading(doc, "六、技术方案", level=1)
    _add_heading(doc, "6.1 项目理解", level=2)
    _add_para(doc, f"本项目属于 {params['industry'] or '相关行业'} 领域，项目需求可概括为：{_project_scope(params['industry'])}")
    _add_heading(doc, "6.2 技术架构", level=2)
    _add_para(doc, "描述总体技术架构、关键技术选型、系统部署方案。")
    _add_heading(doc, "6.3 实施计划", level=2)
    _add_para(doc, f"项目工期/服务期为 {params['duration'] or '按招标文件要求'}，实施计划分为启动、设计、实施、测试、验收、运维等阶段。")
    _add_heading(doc, "6.4 质量与风险管理", level=2)
    _add_para(doc, "说明质量保证措施、风险控制策略、应急预案。")

    _add_heading(doc, "七、商务方案", level=1)
    _add_heading(doc, "7.1 服务承诺", level=2)
    _add_para(doc, "承诺服务内容、响应时间、服务网点、质保期限等。")
    _add_heading(doc, "7.2 项目团队", level=2)
    _add_para(doc, "列示项目经理及关键成员简历、职责分工、联系方式。")

    _add_heading(doc, "八、报价文件", level=1)
    _add_para(doc, f"投标总价为 {params['amount']:,.2f} 元（大写：{__rmb_capital(params['amount'])}），分项报价详见商务报价表。")
    _add_para(doc, "报价有效期、付款方式、税费说明等按招标文件要求填写。")

    _add_heading(doc, "九、售后服务承诺", level=1)
    _add_para(doc, "详细说明售后服务体系、故障响应流程、备件支持、培训计划、定期回访等。")

    _add_heading(doc, "十、附件清单", level=1)
    attachments = [
        "附件 1：营业执照副本复印件",
        "附件 2：法定代表人身份证明及授权书",
        "附件 3：资质证书复印件",
        "附件 4：财务审计报告",
        "附件 5：近 3 年同类项目业绩证明",
        "附件 6：项目经理及关键人员简历",
        "附件 7：商务报价表",
        "附件 8：技术偏离表",
        "附件 9：其他招标文件要求的材料",
    ]
    for att in attachments:
        _add_para(doc, att)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f"{DISCLAIMER} 本投标文件大纲由 prajna 企智自动生成，仅供结构参考，正式投标文件须根据招标文件及企业实际情况调整完善。")
    run.font.name = "微软雅黑"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# Workbook orchestration
# ---------------------------------------------------------------------------
def build_excel(params, output_path):
    wb = Workbook()
    build_cover(wb, params)
    build_tender_breakdown(wb, params)
    build_qualification(wb, params)
    build_tech_deviation(wb, params)
    build_commercial_quote(wb, params)
    build_scoring_index(wb, params)
    build_timeline(wb, params)
    wb.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def parse_args(argv=None):
    parser = argparse.ArgumentParser(description="生成招投标套件（Excel + Word）")
    parser.add_argument("--project", "-p", required=True, help="项目名称")
    parser.add_argument("--bidder", "-b", default="", help="投标人/公司名称")
    parser.add_argument("--tenderer", "-t", default="", help="招标人/采购人")
    parser.add_argument("--amount", "-a", type=float, default=0.0, help="投标总金额（元）")
    parser.add_argument("--duration", "-d", default="", help="工期/服务期，如 180天")
    parser.add_argument("--industry", "-i", default="IT", help="行业领域，如 IT、建筑、服务、制造")
    parser.add_argument("--output", "-o", default=str(DEFAULT_OUTPUT_DIR), help="输出目录或文件路径前缀")
    parser.add_argument("--format", "-f", default="all", choices=["excel", "word", "all"], help="输出格式：excel / word / all，默认 all")
    return parser.parse_args(argv)


def main(argv=None):
    args = parse_args(argv)

    params = {
        "project": args.project,
        "bidder": args.bidder,
        "tenderer": args.tenderer,
        "amount": _parse_amount(args.amount),
        "duration": args.duration,
        "industry": args.industry,
    }

    output = Path(args.output)
    if output.suffix:
        # User gave a file path prefix; derive directory
        output_dir = output.parent
        prefix = output.stem
    else:
        output_dir = output
        prefix = _safe_filename(params["project"])
    output_dir.mkdir(parents=True, exist_ok=True)

    safe_project = _safe_filename(params["project"])
    date_suffix = _today()

    generated = []

    if args.format in ("excel", "all"):
        excel_path = output_dir / f"prajna_招投标套件_{safe_project}_{date_suffix}.xlsx"
        build_excel(params, excel_path)
        generated.append(str(excel_path))
        print(f"已生成 Excel：{excel_path}")

    if args.format in ("word", "all"):
        word_path = output_dir / f"prajna_投标文件大纲_{safe_project}_{date_suffix}.docx"
        build_word_outline(params, word_path)
        generated.append(str(word_path))
        print(f"已生成 Word：{word_path}")

    return 0


if __name__ == "__main__":
    sys.exit(main())
