#!/usr/bin/env python3
"""
prajna-contract-review-assistant 生成器 v1.0.0
一键生成合同审查清单 Word 文档，覆盖商务、法务、财务、技术、风险五大维度。
"""
import argparse
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_DIR = SCRIPT_DIR.parent
SAMPLES_DIR = Path.home() / ".prajna" / "prajna-contract-review-assistant" / "samples"


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_xml = tcPr.find(qn(f"w:{edge}"))
        if edge_xml is None:
            edge_xml = docx.oxml.OxmlElement(f"w:{edge}")
            tcPr.append(edge_xml)
        edge_xml.set(qn("w:val"), kwargs.get("val", "single"))
        edge_xml.set(qn("w:sz"), kwargs.get("sz", "4"))
        edge_xml.set(qn("w:space"), "0")
        edge_xml.set(qn("w:color"), kwargs.get("color", "000000"))


def add_heading_zh(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return p


def add_para(doc, text, bold=False, color=None, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def build_contract_review(args) -> Path:
    doc = Document()
    # Title
    title = doc.add_heading(f"{args.company} 合同审查意见书", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(31, 78, 120)

    add_para(doc, f"合同名称：{args.contract_name}")
    add_para(doc, f"合同类型：{args.contract_type}")
    add_para(doc, f"审查日期：{args.date}")
    add_para(doc, f"审查人：{args.reviewer}")
    doc.add_paragraph()

    # Executive summary
    add_heading_zh(doc, "一、审查结论", level=1)
    add_para(doc, f"经审查，本合同为{args.contract_type}，合同金额 {args.amount:,.0f} 元，合作期限 {args.term}。整体法律风险{'可控' if args.risk_level == '低' else '需重点关注'}，建议在补充以下条款或取得对方确认后签署。")

    # Review dimensions
    dimensions = [
        ("二、主体资格审查", [
            "签约对方是否具备合法存续的营业执照及相应资质；",
            "签约代表是否持有有效授权委托书；",
            "对方是否为失信被执行人或存在重大诉讼。",
        ]),
        ("三、商务条款审查", [
            "标的、数量、质量、价款是否明确且无歧义；",
            "交付时间、地点、方式、验收标准是否清晰；",
            "价格调整机制、税费承担、发票类型是否约定。",
        ]),
        ("四、付款与财务条款", [
            "付款节点与交付/验收是否挂钩；",
            "账期、违约金、滞纳金是否合理；",
            "是否约定预付款比例及对应的履约担保。",
        ]),
        ("五、违约与解除条款", [
            "违约责任是否对等，赔偿上限是否合理；",
            "单方解除权的触发条件是否明确；",
            "不可抗力、情势变更条款是否完整。",
        ]),
        ("六、知识产权与保密", [
            "知识产权归属、许可范围、侵权责任是否明确；",
            "保密信息范围、期限、违约责任是否合理；",
            "数据安全与个人信息处理义务是否约定。",
        ]),
        ("七、争议解决", [
            "争议解决方式（诉讼/仲裁）是否明确；",
            "管辖法院或仲裁机构是否对我方有利；",
            "法律适用是否为中国大陆法律。",
        ]),
    ]

    for heading, items in dimensions:
        add_heading_zh(doc, heading, level=1)
        for item in items:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(item)
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(10.5)

    # Risk register table
    add_heading_zh(doc, "八、风险登记与修改建议", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    headers = ["风险点", "风险等级", "影响", "修改建议", "责任部门"]
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    risks = [
        ["付款节点过早", "高", "资金占用，对方履约动力下降", "建议改为 30%预付+40%到货+30%验收", "财务部"],
        ["验收标准模糊", "中", "易引发质量争议", "补充明确的技术验收指标与测试用例", "技术部"],
        ["违约金上限过低", "中", "违约成本不足以覆盖损失", "建议违约金不低于合同金额 10%", "法务部"],
        ["管辖约定不利", "低", "诉讼成本增加", "优先约定我方所在地法院/仲裁", "法务部"],
    ]
    for row_data in risks:
        row = table.add_row().cells
        for i, text in enumerate(row_data):
            row[i].text = text
            for run in row[i].paragraphs[0].runs:
                run.font.name = "Microsoft YaHei"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                run.font.size = Pt(10)

    add_heading_zh(doc, "九、最终建议", level=1)
    add_para(doc, "1. 建议由业务部门、法务部、财务部联合会签后提交审批；")
    add_para(doc, "2. 对高风险条款应要求对方出具书面补充协议或承诺函；")
    add_para(doc, "3. 签署后应将合同扫描件、审批单、补充协议统一归档至合同管理系统。")

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def parse_args(argv=None):
    parser = argparse.ArgumentParser(description="生成合同审查意见书 Word 文档")
    parser.add_argument("--company", default="智信科技", help="公司名称")
    parser.add_argument("--contract-name", default="软件采购与服务合同", help="合同名称")
    parser.add_argument("--contract-type", default="采购合同", help="合同类型")
    parser.add_argument("--amount", type=float, default=580000, help="合同金额（元）")
    parser.add_argument("--term", default="12个月", help="合作期限")
    parser.add_argument("--risk-level", default="中", choices=["高", "中", "低"], help="风险等级")
    parser.add_argument("--reviewer", default="法务经理", help="审查人")
    parser.add_argument("--date", default=datetime.now().strftime("%Y-%m-%d"), help="审查日期")
    parser.add_argument("--output", "-o", help="输出文件路径")
    return parser.parse_args(argv)


def main(argv=None):
    args = parse_args(argv)
    if not args.output:
        SAMPLES_DIR.mkdir(parents=True, exist_ok=True)
        args.output = SAMPLES_DIR / f"prajna_合同审查意见书_{args.company}_{args.date}.docx"
    path = build_contract_review(args)
    print(f"已生成合同审查意见书：{path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
