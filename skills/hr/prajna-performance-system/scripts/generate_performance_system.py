#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
prajna 企智 - 绩效体系助手生成器 v1.0.0

为企业一键生成完整的绩效管理体系套件：
- Excel 工作簿（6 个工作表）：
  1. KPI 指标库（按部门/岗位分类）
  2. 绩效合同/目标责任书
  3. 绩效评分表（自评/上级评/HR 复核）
  4. 绩效面谈记录
  5. 绩效结果分布
  6. 绩效改进计划 PIP
- Word 文档：绩效管理制度（目的、适用范围、考核周期、流程、结果应用）

依赖：openpyxl、python-docx（未安装时脚本会提示安装命令）
"""

import argparse
import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path

# ---------------------------------------------------------------------------
# Optional dependencies
# ---------------------------------------------------------------------------
try:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    HAVE_OPENPYXL = True
except ImportError:  # pragma: no cover
    HAVE_OPENPYXL = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

    HAVE_DOCX = True
except ImportError:  # pragma: no cover
    HAVE_DOCX = False

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_DIR = SCRIPT_DIR.parent
DATA_DIR = SKILL_DIR / "data"
KPI_LIBRARY_PATH = DATA_DIR / "performance_kpi_library.json"
DEFAULT_SAMPLES_DIR = Path.home() / ".prajna" / "prajna-performance-system" / "samples"

# ---------------------------------------------------------------------------
# Styling constants
# ---------------------------------------------------------------------------
THEME_BLUE = "1F4E78"
THEME_LIGHT_BLUE = "D9E1F2"
THEME_BG = "F2F2F2"
THEME_GREEN = "E2EFDA"
THEME_YELLOW = "FFF2CC"
THEME_RED = "FCE4D6"

HEADER_FILL = PatternFill(start_color=THEME_BLUE, end_color=THEME_BLUE, fill_type="solid")
HEADER_FONT = Font(name="微软雅黑", size=11, bold=True, color="FFFFFF")
SUBHEADER_FILL = PatternFill(start_color=THEME_LIGHT_BLUE, end_color=THEME_LIGHT_BLUE, fill_type="solid")
NORMAL_FONT = Font(name="微软雅黑", size=10)
BOLD_FONT = Font(name="微软雅黑", size=10, bold=True)
SMALL_FONT = Font(name="微软雅黑", size=9)
CENTER_ALIGN = Alignment(horizontal="center", vertical="center", wrap_text=True)
LEFT_ALIGN = Alignment(horizontal="left", vertical="center", wrap_text=True)
THIN_BORDER = Border(
    left=Side(style="thin", color="B7B7B7"),
    right=Side(style="thin", color="B7B7B7"),
    top=Side(style="thin", color="B7B7B7"),
    bottom=Side(style="thin", color="B7B7B7"),
)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def load_json(path: Path) -> dict:
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def _safe_filename(text: str) -> str:
    """移除或替换文件名中的非法字符。"""
    return re.sub(r'[\\/:*?"<>|]+', "_", str(text)).strip().replace(" ", "_") or "default"


def _today() -> str:
    return datetime.now().strftime("%Y-%m-%d")


def _apply_header_row(ws, row_idx, headers, widths=None):
    """给指定行写入表头并设置样式。"""
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=row_idx, column=col_idx, value=header)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = CENTER_ALIGN
        cell.border = THIN_BORDER
    if widths:
        for idx, width in enumerate(widths, 1):
            ws.column_dimensions[get_column_letter(idx)].width = width


def _style_data_cell(cell, font=None, align=None, fill=None):
    cell.font = font or NORMAL_FONT
    cell.alignment = align or LEFT_ALIGN
    cell.border = THIN_BORDER
    if fill:
        cell.fill = fill


def _set_print_area(ws):
    """设置打印区域为已用范围（openpyxl 不支持直接设置，此处预留扩展）。"""
    pass


# ---------------------------------------------------------------------------
# KPI matching
# ---------------------------------------------------------------------------
def match_kpis(library: dict, department: str, position: str):
    """
    根据部门和岗位名称匹配最合适的 KPI 列表。
    先按部门关键词匹配部门，再在部门内按岗位关键词匹配岗位；
    未命中则返回通用岗位模板。
    """
    departments = library.get("departments", {})

    # 1. 精确匹配部门名称
    if department in departments:
        dept_data = departments[department]
        positions = dept_data.get("positions", {})
        if position in positions:
            return department, position, positions[position].get("kpis", [])
        # 2. 在部门内模糊匹配岗位
        for pos_name, pos_data in positions.items():
            if pos_name in position or position in pos_name:
                return department, pos_name, pos_data.get("kpis", [])

    # 3. 按关键词匹配部门
    for dept_name, dept_data in departments.items():
        keywords = dept_data.get("keywords", [])
        if any(kw in department for kw in keywords):
            positions = dept_data.get("positions", {})
            if position in positions:
                return dept_name, position, positions[position].get("kpis", [])
            for pos_name, pos_data in positions.items():
                if pos_name in position or position in pos_name:
                    return dept_name, pos_name, pos_data.get("kpis", [])

    # 4. 全局按岗位关键词匹配
    for dept_name, dept_data in departments.items():
        for pos_name, pos_data in dept_data.get("positions", {}).items():
            if pos_name in position or position in pos_name:
                return dept_name, pos_name, pos_data.get("kpis", [])

    # 5. 回退到通用
    general = departments.get("通用", {}).get("positions", {}).get("通用岗位", {})
    return department or "通用", position or "通用岗位", general.get("kpis", [])


def get_method_info(library: dict, method: str) -> dict:
    methods = library.get("methods", {})
    if method in methods:
        return methods[method]
    # 模糊匹配
    for key, value in methods.items():
        if key in method or method in key:
            return value
    return methods.get("KPI", {"description": "", "steps": []})


def get_level_scheme(library: dict, levels: str) -> list:
    schemes = library.get("level_schemes", {})
    if levels in schemes:
        return schemes[levels].get("levels", [])
    for key, value in schemes.items():
        if key in levels or levels in key:
            return value.get("levels", [])
    return schemes.get("A/B/C/D", {}).get("levels", [])


def normalize_cycle(cycle: str) -> str:
    """统一考核周期表述。"""
    mapping = {
        "month": "月度", "monthly": "月度", "月": "月度",
        "quarter": "季度", "quarterly": "季度", "季": "季度",
        "half": "半年度", "halfyear": "半年度", "半年": "半年度",
        "year": "年度", "yearly": "年度", "年": "年度",
    }
    key = cycle.lower().strip() if cycle else ""
    return mapping.get(key, cycle or "季度")


def normalize_levels(levels: str) -> str:
    """统一绩效等级表述。"""
    text = levels.upper().strip() if levels else ""
    if "/" in text:
        parts = text.split("/")
    else:
        parts = list(text)
    if "S" in parts and "A" in parts:
        return "S/A/B/C/D"
    return "A/B/C/D"


# ---------------------------------------------------------------------------
# Excel builders
# ---------------------------------------------------------------------------
def _add_readme_sheet(wb, meta):
    ws = wb.active
    ws.title = "使用说明"
    ws.append([])
    ws.append(["绩效管理体系套件"])
    ws.append([])
    ws.append(["企业名称", meta["company"]])
    ws.append(["考核部门", meta["department"]])
    ws.append(["考核岗位", meta["position"]])
    ws.append(["考核周期", meta["cycle"]])
    ws.append(["考核方法", meta["method"]])
    ws.append(["绩效等级", meta["levels"]])
    ws.append(["生成日期", meta["date"]])
    ws.append([])
    ws.append(["本 Excel 包含六张工作表："])
    ws.append(["1. KPI 指标库", "按部门/岗位分类的核心 KPI 指标、权重、目标值、计算方式、数据来源与评分规则。"])
    ws.append(["2. 绩效合同", "员工绩效目标责任书模板，含目标确认、签字栏与双方承诺。"])
    ws.append(["3. 绩效评分表", "自评、上级评、HR 复核三栏评分表，支持自动计算加权得分与总得分。"])
    ws.append(["4. 绩效面谈记录", "绩效面谈结构化记录表，含面谈准备、关键对话与后续行动。"])
    ws.append(["5. 绩效结果分布", "团队绩效等级分布与比例指引，便于校准与人才盘点。"])
    ws.append(["6. PIP 改进计划", "绩效改进计划模板，用于 D 级或连续低绩效员工。"])
    ws.append([])
    ws.append(["填写须知"])
    ws.append(["1.", "KPI 目标值为行业参考，企业应结合自身战略、历史数据与岗位实际调整。"])
    ws.append(["2.", "权重合计应等于 100%，生成时已做校验；如需增删指标，请重新分配权重。"])
    ws.append(["3.", "绩效评分表中的「单项得分」为百分制，按 KPI 指标库中的评分规则评定。"])
    ws.append(["4.", "绩效结果分布中的比例为指导性建议，强制分布需经民主程序并向员工公示。"])
    ws.append([])
    ws.append(["免责声明"])
    ws.append(["本套件由 prajna 企智人工智能生成，仅供业务参考，不构成正式绩效考核或劳动用工依据，使用前请由企业 HR、法务及管理层审核确认。"])
    ws.append([])
    ws.append(["与现有 Skill 的打通"])
    ws.append(["- KPI 指标库复用了 prajna-salary-template、prajna-clothing-teamleader-duty 的岗位 KPI 思路。"])
    ws.append(["- 绩效结果可与 prajna-salary-template / prajna-compensation-system 联动，作为奖金、调薪依据。"])

    # 合并标题行
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=2)
    title_cell = ws.cell(row=2, column=1)
    title_cell.font = Font(name="微软雅黑", size=16, bold=True, color=THEME_BLUE)
    title_cell.alignment = CENTER_ALIGN

    for row in ws.iter_rows(min_row=4, max_row=10, min_col=1, max_col=2):
        row[0].font = BOLD_FONT
        row[0].alignment = LEFT_ALIGN
        row[1].alignment = LEFT_ALIGN

    ws.column_dimensions["A"].width = 22
    ws.column_dimensions["B"].width = 95


def _add_kpi_library_sheet(wb, kpis, meta):
    ws = wb.create_sheet("KPI指标库")
    headers = ["KPI编号", "KPI名称", "权重(%)", "目标值", "计算方式", "数据来源", "评分规则", "考核周期", "适用岗位"]
    widths = [12, 18, 10, 14, 28, 18, 34, 10, 16]
    _apply_header_row(ws, 1, headers, widths)
    ws.freeze_panes = "A2"

    total_weight = 0
    for row_idx, kpi in enumerate(kpis, 2):
        weight = kpi.get("weight", 0)
        total_weight += weight
        values = [
            kpi.get("id", ""),
            kpi.get("name", ""),
            weight,
            kpi.get("target", ""),
            kpi.get("formula", ""),
            kpi.get("source", ""),
            kpi.get("rule", ""),
            kpi.get("cycle", meta["cycle"]),
            meta["position"],
        ]
        for col_idx, value in enumerate(values, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            align = CENTER_ALIGN if col_idx in (1, 3, 4, 8, 9) else LEFT_ALIGN
            _style_data_cell(cell, align=align)
        ws.row_dimensions[row_idx].height = 42

    # 权重合计校验行
    total_row = len(kpis) + 2
    ws.cell(row=total_row, column=1, value="合计")
    ws.cell(row=total_row, column=2, value="权重合计应=100%")
    weight_cell = ws.cell(row=total_row, column=3, value=total_weight)
    weight_cell.font = BOLD_FONT
    weight_cell.alignment = CENTER_ALIGN
    weight_cell.fill = PatternFill(start_color=THEME_LIGHT_BLUE, end_color=THEME_LIGHT_BLUE, fill_type="solid")
    for col in range(1, 10):
        cell = ws.cell(row=total_row, column=col)
        cell.border = THIN_BORDER
        cell.font = BOLD_FONT
        cell.alignment = CENTER_ALIGN
    ws.cell(row=total_row, column=3).border = THIN_BORDER

    # 在末尾添加免责声明行
    note_row = total_row + 2
    ws.cell(row=note_row, column=1, value="说明")
    ws.cell(row=note_row, column=1).font = BOLD_FONT
    ws.merge_cells(start_row=note_row, start_column=2, end_row=note_row, end_column=9)
    note_cell = ws.cell(row=note_row, column=2, value="本页 KPI 为行业参考模板，企业应根据实际业务目标、数据可获取性与岗位职责进行调整。")
    note_cell.font = SMALL_FONT
    note_cell.alignment = LEFT_ALIGN


def _add_contract_sheet(wb, kpis, meta):
    ws = wb.create_sheet("绩效合同")
    row = 1
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=7)
    title = ws.cell(row=row, column=1, value=f"{meta['company']} {meta['cycle']}绩效目标责任书")
    title.font = Font(name="微软雅黑", size=14, bold=True, color=THEME_BLUE)
    title.alignment = CENTER_ALIGN
    ws.row_dimensions[row].height = 28
    row += 2

    info = [
        ["责任部门", meta["department"], "岗位名称", meta["position"], "考核周期", meta["cycle"], ""],
        ["考核方法", meta["method"], "签订日期", meta["date"], "", "", ""],
    ]
    for r in info:
        for col_idx, value in enumerate(r, 1):
            cell = ws.cell(row=row, column=col_idx, value=value)
            if col_idx in (1, 3, 5):
                cell.font = BOLD_FONT
                cell.fill = PatternFill(start_color=THEME_BG, end_color=THEME_BG, fill_type="solid")
            cell.alignment = CENTER_ALIGN if col_idx in (1, 3, 5) else LEFT_ALIGN
            cell.border = THIN_BORDER
        row += 1

    row += 1
    ws.cell(row=row, column=1, value="一、考核目标").font = BOLD_FONT
    row += 1

    headers = ["序号", "KPI编号", "KPI名称", "权重(%)", "目标值", "数据来源", "考核周期"]
    widths = [8, 12, 18, 10, 14, 18, 10]
    _apply_header_row(ws, row, headers, widths)
    header_row = row
    row += 1

    for idx, kpi in enumerate(kpis, 1):
        values = [idx, kpi.get("id", ""), kpi.get("name", ""), kpi.get("weight", 0), kpi.get("target", ""), kpi.get("source", ""), kpi.get("cycle", meta["cycle"])]
        for col_idx, value in enumerate(values, 1):
            cell = ws.cell(row=row, column=col_idx, value=value)
            align = CENTER_ALIGN if col_idx in (1, 2, 4, 5, 7) else LEFT_ALIGN
            _style_data_cell(cell, align=align)
        ws.row_dimensions[row].height = 32
        row += 1

    row += 1
    ws.cell(row=row, column=1, value="二、绩效等级与结果应用").font = BOLD_FONT
    row += 1
    level_headers = ["等级", "分数区间", "结果应用"]
    _apply_header_row(ws, row, level_headers)
    ws.column_dimensions["A"].width = 12
    ws.column_dimensions["B"].width = 18
    ws.column_dimensions["C"].width = 60
    level_row = row
    row += 1
    for lvl in meta["level_scheme"]:
        values = [lvl["level"], f"{lvl['score_min']}-{lvl['score_max']}", lvl["application"]]
        for col_idx, value in enumerate(values, 1):
            cell = ws.cell(row=row, column=col_idx, value=value)
            align = CENTER_ALIGN if col_idx in (1, 2) else LEFT_ALIGN
            _style_data_cell(cell, align=align)
        row += 1

    row += 1
    ws.cell(row=row, column=1, value="三、双方承诺").font = BOLD_FONT
    row += 1
    commitment_text = (
        "1. 员工承诺：本人已充分理解上述绩效目标，将全力以赴完成各项考核指标，"
        "定期向上级汇报进度，主动寻求资源支持。\n"
        "2. 上级承诺：本人将与员工保持持续绩效沟通，提供必要的辅导、资源与反馈，"
        "确保考核过程公平、公正、公开。\n"
        "3. 双方确认：本责任书经双方签字后生效，作为绩效评分、奖金发放与人才发展的重要依据。"
    )
    ws.merge_cells(start_row=row, start_column=1, end_row=row + 2, end_column=7)
    cell = ws.cell(row=row, column=1, value=commitment_text)
    cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
    cell.border = THIN_BORDER
    cell.font = NORMAL_FONT
    row += 3

    row += 1
    sign_row = row
    ws.merge_cells(start_row=sign_row, start_column=1, end_row=sign_row, end_column=3)
    ws.merge_cells(start_row=sign_row, start_column=4, end_row=sign_row, end_column=7)
    ws.cell(row=sign_row, column=1, value="员工签字：                日期：").alignment = LEFT_ALIGN
    ws.cell(row=sign_row, column=4, value="直接上级签字：                日期：").alignment = LEFT_ALIGN
    ws.row_dimensions[sign_row].height = 36


def _add_score_sheet(wb, kpis, meta):
    ws = wb.create_sheet("绩效评分表")
    row = 1
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=11)
    title = ws.cell(row=row, column=1, value=f"{meta['cycle']}绩效评分表（{meta['department']} - {meta['position']}）")
    title.font = Font(name="微软雅黑", size=14, bold=True, color=THEME_BLUE)
    title.alignment = CENTER_ALIGN
    ws.row_dimensions[row].height = 26
    row += 2

    info = [
        ["员工姓名", "", "部门", meta["department"], "岗位", meta["position"], "考核周期", meta["cycle"]],
        ["直接上级", "", "HR 复核人", "", "评分日期", meta["date"], "", ""],
    ]
    for r in info:
        for col_idx, value in enumerate(r, 1):
            cell = ws.cell(row=row, column=col_idx, value=value)
            if col_idx in (1, 3, 5, 7):
                cell.font = BOLD_FONT
                cell.fill = PatternFill(start_color=THEME_BG, end_color=THEME_BG, fill_type="solid")
            cell.alignment = CENTER_ALIGN
            cell.border = THIN_BORDER
        row += 1

    row += 1
    headers = ["KPI编号", "KPI名称", "权重(%)", "目标值", "实际完成", "自评得分", "上级评分", "HR复核分", "最终得分", "加权得分", "备注"]
    widths = [12, 18, 10, 14, 14, 12, 12, 12, 12, 12, 20]
    _apply_header_row(ws, row, headers, widths)
    header_row = row
    ws.freeze_panes = f"A{header_row + 1}"
    row += 1

    for kpi in kpis:
        values = [
            kpi.get("id", ""),
            kpi.get("name", ""),
            kpi.get("weight", 0),
            kpi.get("target", ""),
            "",  # 实际完成
            "",  # 自评得分
            "",  # 上级评分
            "",  # HR复核分
            "",  # 最终得分
            "",  # 加权得分
            "",  # 备注
        ]
        for col_idx, value in enumerate(values, 1):
            cell = ws.cell(row=row, column=col_idx, value=value)
            align = CENTER_ALIGN if col_idx in (1, 3, 4, 5, 6, 7, 8, 9, 10) else LEFT_ALIGN
            _style_data_cell(cell, align=align)
        # 最终得分 = 自评*0.2 + 上级*0.5 + HR复核*0.3
        # 加权得分 = 最终得分 * 权重 / 100
        formula_final = f'=IF(COUNTA(F{row}:H{row})=3,ROUND(F{row}*0.2+G{row}*0.5+H{row}*0.3,2),"")'
        ws.cell(row=row, column=9, value=formula_final)
        formula_weighted = f'=IF(AND(C{row}<>"",I{row}<>""),ROUND(C{row}*I{row}/100,2),"")'
        ws.cell(row=row, column=10, value=formula_weighted)
        ws.row_dimensions[row].height = 34
        row += 1

    # 汇总行
    total_row = row
    ws.cell(row=total_row, column=1, value="合计")
    ws.cell(row=total_row, column=2, value="绩效总得分")
    for col in range(1, 10):
        cell = ws.cell(row=total_row, column=col)
        cell.font = BOLD_FONT
        cell.alignment = CENTER_ALIGN
        cell.border = THIN_BORDER
    total_formula = f'=IF(COUNTA(J{header_row + 1}:J{total_row - 1})>0,SUM(J{header_row + 1}:J{total_row - 1}),"")'
    total_cell = ws.cell(row=total_row, column=10, value=total_formula)
    total_cell.font = BOLD_FONT
    total_cell.alignment = CENTER_ALIGN
    total_cell.border = THIN_BORDER
    total_cell.fill = PatternFill(start_color=THEME_LIGHT_BLUE, end_color=THEME_LIGHT_BLUE, fill_type="solid")
    ws.cell(row=total_row, column=11).border = THIN_BORDER

    row += 2
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=11)
    note = ws.cell(
        row=row,
        column=1,
        value="评分说明：自评占 20%、上级评分占 50%、HR 复核占 30%；企业可根据管理需要调整权重。最终得分 90-100 为优秀，80-89 为良好，60-79 为合格，60 以下为不合格。"
    )
    note.font = SMALL_FONT
    note.alignment = LEFT_ALIGN


def _add_interview_sheet(wb, meta):
    ws = wb.create_sheet("绩效面谈记录")
    row = 1
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
    title = ws.cell(row=row, column=1, value="绩效面谈记录表")
    title.font = Font(name="微软雅黑", size=14, bold=True, color=THEME_BLUE)
    title.alignment = CENTER_ALIGN
    ws.row_dimensions[row].height = 26
    row += 2

    basic = [
        ["员工姓名", "", "部门", meta["department"]],
        ["岗位名称", meta["position"], "考核周期", meta["cycle"]],
        ["面谈日期", meta["date"], "面谈地点", ""],
        ["直接上级", "", "HR 列席", ""],
    ]
    for r in basic:
        for col_idx, value in enumerate(r, 1):
            cell = ws.cell(row=row, column=col_idx, value=value)
            if col_idx in (1, 3):
                cell.font = BOLD_FONT
                cell.fill = PatternFill(start_color=THEME_BG, end_color=THEME_BG, fill_type="solid")
            cell.alignment = CENTER_ALIGN
            cell.border = THIN_BORDER
        row += 1

    sections = [
        ("一、绩效回顾（员工自评）", "请员工简要总结本周期主要工作成果、亮点与不足。"),
        ("二、上级反馈", "上级对员工目标完成情况、关键行为与能力的评价。"),
        ("三、发展需求", "员工提出的能力提升、资源支持或职业发展需求。"),
        ("四、下期目标共识", "双方确认的下一周期核心目标与关键行动。"),
        ("五、后续行动计划", "明确的改进措施、责任人、完成时间与验收标准。"),
        ("六、员工申诉/备注", "员工对考核结果的异议或补充说明。"),
    ]

    for section_title, placeholder in sections:
        row += 1
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
        header = ws.cell(row=row, column=1, value=section_title)
        header.font = BOLD_FONT
        header.fill = SUBHEADER_FILL
        header.alignment = LEFT_ALIGN
        header.border = THIN_BORDER
        ws.row_dimensions[row].height = 24
        row += 1
        ws.merge_cells(start_row=row, start_column=1, end_row=row + 2, end_column=4)
        cell = ws.cell(row=row, column=1, value=placeholder)
        cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
        cell.border = THIN_BORDER
        cell.font = NORMAL_FONT
        ws.row_dimensions[row].height = 60
        ws.row_dimensions[row + 1].height = 20
        ws.row_dimensions[row + 2].height = 20
        row += 3

    row += 1
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=2)
    ws.merge_cells(start_row=row, start_column=3, end_row=row, end_column=4)
    ws.cell(row=row, column=1, value="员工签字：                日期：").alignment = LEFT_ALIGN
    ws.cell(row=row, column=3, value="上级签字：                日期：").alignment = LEFT_ALIGN
    ws.row_dimensions[row].height = 36

    ws.column_dimensions["A"].width = 16
    ws.column_dimensions["B"].width = 24
    ws.column_dimensions["C"].width = 16
    ws.column_dimensions["D"].width = 50


def _add_distribution_sheet(wb, meta):
    ws = wb.create_sheet("绩效结果分布")
    row = 1
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
    title = ws.cell(row=row, column=1, value=f"{meta['cycle']}绩效结果分布与比例指引")
    title.font = Font(name="微软雅黑", size=14, bold=True, color=THEME_BLUE)
    title.alignment = CENTER_ALIGN
    ws.row_dimensions[row].height = 26
    row += 2

    headers = ["绩效等级", "分数区间", "等级名称", "建议比例", "结果应用", "实际人数（手动填写）"]
    widths = [12, 16, 12, 14, 40, 20]
    _apply_header_row(ws, row, headers, widths)
    header_row = row
    row += 1

    colors = {
        "S": THEME_GREEN,
        "A": THEME_GREEN,
        "B": THEME_YELLOW,
        "C": THEME_YELLOW,
        "D": THEME_RED,
    }

    for lvl in meta["level_scheme"]:
        values = [
            lvl["level"],
            f"{lvl['score_min']}-{lvl['score_max']}",
            lvl["name"],
            lvl["ratio_guidance"],
            lvl["application"],
            "",
        ]
        fill = PatternFill(start_color=colors.get(lvl["level"], "FFFFFF"), end_color=colors.get(lvl["level"], "FFFFFF"), fill_type="solid")
        for col_idx, value in enumerate(values, 1):
            cell = ws.cell(row=row, column=col_idx, value=value)
            align = CENTER_ALIGN if col_idx in (1, 2, 3, 4, 6) else LEFT_ALIGN
            _style_data_cell(cell, align=align, fill=fill)
        row += 1

    row += 1
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
    note = ws.cell(
        row=row,
        column=1,
        value="说明：建议比例为参考值，实际分布应结合团队业绩、目标难度与文化导向综合确定；强制分布建议经民主程序并向员工公示。"
    )
    note.font = SMALL_FONT
    note.alignment = LEFT_ALIGN

    row += 2
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
    ws.cell(row=row, column=1, value="团队绩效分布统计示例（按 20 人团队估算）").font = BOLD_FONT
    row += 1
    sample_headers = ["等级", "建议比例", "估算人数", "实际人数", "偏差说明", ""]
    _apply_header_row(ws, row, sample_headers)
    row += 1
    for lvl in meta["level_scheme"]:
        ratio_text = lvl["ratio_guidance"].replace("≤", "").replace("≈", "").replace("%", "")
        try:
            ratio = float(ratio_text) / 100
        except ValueError:
            ratio = 0
        estimate = round(20 * ratio)
        values = [lvl["level"], lvl["ratio_guidance"], estimate, "", "", ""]
        for col_idx, value in enumerate(values, 1):
            cell = ws.cell(row=row, column=col_idx, value=value)
            align = CENTER_ALIGN if col_idx in (1, 2, 3, 4) else LEFT_ALIGN
            _style_data_cell(cell, align=align)
        row += 1


def _add_pip_sheet(wb, meta):
    ws = wb.create_sheet("PIP改进计划")
    row = 1
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=5)
    title = ws.cell(row=row, column=1, value="绩效改进计划（PIP）")
    title.font = Font(name="微软雅黑", size=14, bold=True, color=THEME_BLUE)
    title.alignment = CENTER_ALIGN
    ws.row_dimensions[row].height = 26
    row += 2

    basic = [
        ["员工姓名", "", "部门", meta["department"], "岗位", meta["position"]],
        ["绩效等级", "D / 连续 C", "考核周期", meta["cycle"], "启动日期", meta["date"]],
        ["直接上级", "", "HR 负责人", "", "计划周期", "30-90天"],
    ]
    for r in basic:
        for col_idx, value in enumerate(r, 1):
            cell = ws.cell(row=row, column=col_idx, value=value)
            if col_idx in (1, 3, 5):
                cell.font = BOLD_FONT
                cell.fill = PatternFill(start_color=THEME_BG, end_color=THEME_BG, fill_type="solid")
            cell.alignment = CENTER_ALIGN
            cell.border = THIN_BORDER
        row += 1

    row += 1
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=5)
    ws.cell(row=row, column=1, value="一、待改进项").font = BOLD_FONT
    row += 1
    headers = ["序号", "待改进指标/行为", "当前表现", "目标标准", "衡量方式"]
    widths = [8, 26, 26, 26, 26]
    _apply_header_row(ws, row, headers, widths)
    row += 1
    for i in range(1, 5):
        for col_idx in range(1, 6):
            cell = ws.cell(row=row, column=col_idx, value="" if col_idx > 1 else i)
            align = CENTER_ALIGN if col_idx == 1 else LEFT_ALIGN
            _style_data_cell(cell, align=align)
        row += 1

    row += 1
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=5)
    ws.cell(row=row, column=1, value="二、改进行动计划").font = BOLD_FONT
    row += 1
    action_headers = ["序号", "改进行动", "责任人", "完成时间", "验收标准"]
    _apply_header_row(ws, row, action_headers)
    row += 1
    for i in range(1, 5):
        for col_idx in range(1, 6):
            cell = ws.cell(row=row, column=col_idx, value="" if col_idx > 1 else i)
            align = CENTER_ALIGN if col_idx == 1 else LEFT_ALIGN
            _style_data_cell(cell, align=align)
        row += 1

    row += 1
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=5)
    ws.cell(row=row, column=1, value="三、辅导与资源支持").font = BOLD_FONT
    row += 1
    support_items = [
        "1. 上级辅导频率：每周至少 1 次一对一沟通，记录沟通要点。",
        "2. 培训资源：安排岗位技能培训（如业务知识、工具使用、流程规范）。",
        "3. 资源支持：明确所需工具、数据、权限或协作方支持。",
        "4. 里程碑检查：按周/双周检查进展，及时调整行动计划。",
    ]
    for item in support_items:
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=5)
        cell = ws.cell(row=row, column=1, value=item)
        cell.alignment = LEFT_ALIGN
        cell.border = THIN_BORDER
        cell.font = NORMAL_FONT
        row += 1

    row += 1
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=5)
    ws.cell(row=row, column=1, value="四、结果约定").font = BOLD_FONT
    row += 1
    outcomes = [
        "1. 达成目标：员工在 PIP 周期内达成约定目标，视为改进成功，可恢复正常绩效流程。",
        "2. 部分达成：根据改善程度决定是否延长 PIP 周期或调整岗位。",
        "3. 未达成：PIP 周期结束后仍未达标的，按公司制度进入调岗、降薪或依法解除劳动关系程序。",
    ]
    for item in outcomes:
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=5)
        cell = ws.cell(row=row, column=1, value=item)
        cell.alignment = LEFT_ALIGN
        cell.border = THIN_BORDER
        cell.font = NORMAL_FONT
        row += 1

    row += 1
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=2)
    ws.merge_cells(start_row=row, start_column=3, end_row=row, end_column=5)
    ws.cell(row=row, column=1, value="员工签字：                日期：").alignment = LEFT_ALIGN
    ws.cell(row=row, column=3, value="上级与 HR 签字：                日期：").alignment = LEFT_ALIGN
    ws.row_dimensions[row].height = 36

    ws.column_dimensions["A"].width = 10
    ws.column_dimensions["B"].width = 28
    ws.column_dimensions["C"].width = 20
    ws.column_dimensions["D"].width = 20
    ws.column_dimensions["E"].width = 28


def generate_excel(output_path: Path, meta, kpis):
    if not HAVE_OPENPYXL:
        print("[错误] 缺少 openpyxl，请执行：pip install openpyxl", file=sys.stderr)
        sys.exit(1)

    wb = Workbook()
    _add_readme_sheet(wb, meta)
    _add_kpi_library_sheet(wb, kpis, meta)
    _add_contract_sheet(wb, kpis, meta)
    _add_score_sheet(wb, kpis, meta)
    _add_interview_sheet(wb, meta)
    _add_distribution_sheet(wb, meta)
    _add_pip_sheet(wb, meta)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# Word builder
# ---------------------------------------------------------------------------
def _add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def _add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "微软雅黑"
    run.bold = bold
    return p


def _add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "微软雅黑"
    return p


def _add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "微软雅黑"
    return p


def generate_word(output_path: Path, meta, kpis):
    if not HAVE_DOCX:
        print("[错误] 缺少 python-docx，请执行：pip install python-docx", file=sys.stderr)
        sys.exit(1)

    doc = Document()

    # Title
    title = doc.add_heading(f"{meta['company']}绩效管理制度", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = "微软雅黑"
        run.font.size = Pt(18)
        run.font.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"考核方法：{meta['method']} | 考核周期：{meta['cycle']} | 版本：V1.0")
    run.font.size = Pt(10)
    run.font.name = "微软雅黑"
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    # 1. 目的
    _add_heading(doc, "一、目的", level=1)
    _add_paragraph(doc, "为建立科学、公正、透明的绩效管理体系，持续激发员工潜能，支撑公司战略目标的达成，特制定本制度。")
    _add_paragraph(doc, "本制度旨在：")
    _add_bullet(doc, "明确组织目标向个人目标的分解路径，确保上下同欲。")
    _add_bullet(doc, "为员工提供清晰的工作标准、反馈机制与发展指引。")
    _add_bullet(doc, "为薪酬调整、奖金分配、晋升发展、培训培养、人员优化提供客观依据。")

    # 2. 适用范围
    _add_heading(doc, "二、适用范围", level=1)
    _add_paragraph(doc, f"本制度适用于 {meta['company']} 全体正式员工。试用期员工、实习生及外包人员可参照执行，具体由人力资源部另行规定。")

    # 3. 考核原则
    _add_heading(doc, "三、考核原则", level=1)
    _add_numbered(doc, "战略导向：绩效目标应承接公司年度经营计划与部门关键任务。")
    _add_numbered(doc, "SMART 原则：目标应具体（Specific）、可衡量（Measurable）、可达成（Achievable）、相关（Relevant）、有时限（Time-bound）。")
    _add_numbered(doc, "公平公正：考核标准公开透明，评分依据客观数据与事实，允许员工申诉。")
    _add_numbered(doc, "持续沟通：管理者应与员工保持目标设定、过程辅导、结果反馈的闭环沟通。")
    _add_numbered(doc, "结果应用：绩效结果与薪酬、奖金、晋升、培训、淘汰机制紧密联动。")

    # 4. 考核周期
    _add_heading(doc, "四、考核周期", level=1)
    _add_paragraph(doc, f"本制度采用「{meta['cycle']}考核」为主、年度汇总为辅的考核模式。")
    _add_bullet(doc, "目标设定：考核周期初 5 个工作日内完成。")
    _add_bullet(doc, "过程跟踪：管理者每月/每季度与员工进行至少 1 次绩效沟通。")
    _add_bullet(doc, "期末评估：考核周期结束后 10 个工作日内完成自评、上级评与 HR 复核。")
    _add_bullet(doc, "结果审批：部门负责人与 HR 共同审核绩效结果，报管理层批准后生效。")

    # 5. 考核方法
    _add_heading(doc, "五、考核方法", level=1)
    method_info = meta.get("method_info", {})
    _add_paragraph(doc, f"本制度主要采用 {meta['method']} 方法：{method_info.get('description', '通过关键绩效指标量化评估员工目标达成情况。')}")
    if method_info.get("steps"):
        _add_paragraph(doc, "实施步骤：")
        for step in method_info["steps"]:
            _add_bullet(doc, step)

    # 6. 考核流程
    _add_heading(doc, "六、考核流程", level=1)
    _add_numbered(doc, "目标设定（Plan）：上级与员工共同制定绩效合同，明确 KPI、权重、目标值与数据来源。")
    _add_numbered(doc, "过程辅导（Do & Check）：上级持续跟踪进度，提供资源支持与及时反馈，记录关键事件。")
    _add_numbered(doc, "期末评估（Review）：员工自评、上级评分、HR 复核，必要时引入 360 评估。")
    _add_numbered(doc, "结果校准（Calibrate）：部门内进行绩效结果分布校准，避免评分宽松或严格偏差。")
    _add_numbered(doc, "反馈面谈（Feedback）：上级与员工进行一对一面谈，确认结果、共识发展计划。")
    _add_numbered(doc, "结果应用（Apply）：根据绩效等级实施奖金、调薪、晋升、培训或 PIP。")

    # 7. 绩效等级
    _add_heading(doc, "七、绩效等级与结果应用", level=1)
    _add_paragraph(doc, f"绩效等级划分为：{meta['levels']}。具体标准如下：")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "等级"
    hdr_cells[1].text = "分数区间"
    hdr_cells[2].text = "建议比例"
    hdr_cells[3].text = "结果应用"
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = "微软雅黑"

    for lvl in meta["level_scheme"]:
        row_cells = table.add_row().cells
        row_cells[0].text = f"{lvl['level']}（{lvl['name']}）"
        row_cells[1].text = f"{lvl['score_min']}-{lvl['score_max']}"
        row_cells[2].text = lvl["ratio_guidance"]
        row_cells[3].text = lvl["application"]
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "微软雅黑"

    _add_paragraph(doc, "")
    _add_paragraph(doc, "绩效结果应用说明：", bold=True)
    _add_bullet(doc, "薪酬与奖金：绩效等级直接影响绩效奖金系数、季度/年度奖金分配。")
    _add_bullet(doc, "调薪：连续 A 级或 S 级员工优先获得晋升与调薪机会。")
    _add_bullet(doc, "培训发展：B/C 级员工可获得针对性培训，D 级员工进入 PIP。")
    _add_bullet(doc, "岗位调整：PIP 期满仍未改进者，按公司制度进行调岗或依法处理。")

    # 8. 绩效申诉
    _add_heading(doc, "八、绩效申诉", level=1)
    _add_paragraph(doc, "员工如对绩效结果有异议，可在结果公布后 5 个工作日内向人力资源部提交书面申诉。HR 应在 10 个工作日内组织复核并反馈最终结论。")

    # 9. 附则
    _add_heading(doc, "九、附则", level=1)
    _add_numbered(doc, "本制度由人力资源部负责解释与修订。")
    _add_numbered(doc, "本制度自发布之日起生效，原有相关制度同时废止。")
    _add_numbered(doc, "各部门可根据本制度制定实施细则，报人力资源部备案。")

    # Disclaimer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("【人工智能生成-需人工核验】本制度由 prajna 企智 prajna-performance-system 技能根据通用管理实践生成，仅供企业参考。最终制度文本须由企业 HR、法务及管理层结合当地法律法规与公司实际审核确认。")
    run.font.size = Pt(9)
    run.font.name = "微软雅黑"
    run.font.color.rgb = RGBColor(128, 128, 128)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def parse_args(argv=None):
    parser = argparse.ArgumentParser(description="生成企业绩效管理体系套件（Excel + Word）")
    parser.add_argument("--department", "-d", default="电商运营部", help="考核部门名称，默认 电商运营部")
    parser.add_argument("--position", "-p", default="电商运营助理", help="考核岗位名称，默认 电商运营助理")
    parser.add_argument("--cycle", "-c", default="季度", help="考核周期，如 月度、季度、半年度、年度，默认 季度")
    parser.add_argument("--method", "-m", default="KPI", help="考核方法，如 KPI、OKR、360、MBO，默认 KPI")
    parser.add_argument("--levels", "-l", default="A/B/C/D", help="绩效等级划分，如 A/B/C/D 或 S/A/B/C/D，默认 A/B/C/D")
    parser.add_argument("--company", "-co", default="prajna示范企业", help="企业名称")
    parser.add_argument("--output", "-o", help="输出目录或完整路径（覆盖默认输出目录）")
    parser.add_argument(
        "--format",
        "-f",
        choices=["excel", "word", "all"],
        default="all",
        help="输出格式：excel、word、all，默认 all",
    )
    return parser.parse_args(argv)


def build_metadata(args, library):
    cycle = normalize_cycle(args.cycle)
    levels = normalize_levels(args.levels)
    level_scheme = get_level_scheme(library, levels)
    method_info = get_method_info(library, args.method)
    return {
        "company": args.company,
        "department": args.department,
        "position": args.position,
        "cycle": cycle,
        "method": args.method.upper(),
        "levels": levels,
        "level_scheme": level_scheme,
        "method_info": method_info,
        "date": _today(),
    }


def main(argv=None):
    args = parse_args(argv)

    if not HAVE_OPENPYXL:
        print("[错误] 缺少 openpyxl，请执行：pip install openpyxl", file=sys.stderr)
        return 1
    if not HAVE_DOCX and args.format in ("word", "all"):
        print("[错误] 缺少 python-docx，请执行：pip install python-docx", file=sys.stderr)
        return 1

    if not KPI_LIBRARY_PATH.exists():
        print(f"[错误] KPI 库文件不存在：{KPI_LIBRARY_PATH}", file=sys.stderr)
        return 1

    library = load_json(KPI_LIBRARY_PATH)
    meta = build_metadata(args, library)
    matched_dept, matched_pos, kpis = match_kpis(library, args.department, args.position)

    if not kpis:
        print(f"[警告] 未匹配到 {args.department}/{args.position} 的 KPI 模板，将使用通用模板。", file=sys.stderr)
        _, _, kpis = match_kpis(library, "通用", "通用岗位")

    print(f"已匹配 KPI 模板：{matched_dept} / {matched_pos}，共 {len(kpis)} 项指标")

    # Determine output paths
    company_part = _safe_filename(meta["company"])
    dept_part = _safe_filename(meta["department"])
    pos_part = _safe_filename(meta["position"])
    date_part = meta["date"].replace("-", "")

    if args.output:
        output_base = Path(args.output).expanduser().resolve()
        if output_base.suffix:
            # Treat as file path; derive directory and stem
            output_dir = output_base.parent
            stem = output_base.stem
        else:
            output_dir = output_base
            stem = f"prajna_绩效体系_{company_part}_{dept_part}_{pos_part}_{date_part}"
    else:
        output_dir = DEFAULT_SAMPLES_DIR
        stem = f"prajna_绩效体系_{company_part}_{dept_part}_{pos_part}_{date_part}"

    output_dir.mkdir(parents=True, exist_ok=True)

    generated = []
    if args.format in ("excel", "all"):
        excel_path = output_dir / f"{stem}.xlsx"
        generate_excel(excel_path, meta, kpis)
        generated.append(str(excel_path))
        print(f"已生成 Excel：{excel_path}")

    if args.format in ("word", "all"):
        word_path = output_dir / f"{stem}.docx"
        generate_word(word_path, meta, kpis)
        generated.append(str(word_path))
        print(f"已生成 Word：{word_path}")

    return 0


if __name__ == "__main__":
    sys.exit(main())
