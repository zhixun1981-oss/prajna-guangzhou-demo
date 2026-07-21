#!/usr/bin/env python3
"""
prajna-recruitment-assistant 生成器 v1.0.0
为招聘岗位一键生成完整招聘套件：
- Excel：岗位说明书、任职资格与胜任力模型、结构化面试评估表、招聘漏斗与周报、人才画像卡、Offer 薪资建议
- Word：完整招聘套件（JD + 面试提纲 + Offer 模板）

与城市薪酬系数、职级带宽、薪资模板保持数据一致。
"""

import argparse
import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

try:
    import docx
    import docx.oxml.ns
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    HAVE_DOCX = True
except ImportError:
    HAVE_DOCX = False


# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_DIR = SCRIPT_DIR.parent
DATA_DIR = SKILL_DIR / "data"
DEFAULT_SAMPLES_DIR = Path.home() / ".prajna" / "prajna-recruitment-assistant" / "samples"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def load_json(path):
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def _safe_filename(name):
    name = re.sub(r"[\\/:*?\"<>|\s]+", "_", str(name))
    return name.strip("_") or "position"


def _base_styles():
    return {
        "title_font": Font(name="Arial", bold=True, size=16, color="1F4E78"),
        "section_font": Font(name="Arial", bold=True, size=12, color="1F4E78"),
        "header_font": Font(name="Arial", bold=True, color="FFFFFF", size=11),
        "header_fill": PatternFill("solid", fgColor="1F4E78"),
        "subheader_fill": PatternFill("solid", fgColor="4472C4"),
        "blue_font": Font(name="Arial", color="0000FF", size=10),
        "black_font": Font(name="Arial", color="000000", size=10),
        "gray_font": Font(name="Arial", color="808080", size=10),
        "orange_font": Font(name="Arial", color="FF6600", size=10),
        "center_align": Alignment(horizontal="center", vertical="center", wrap_text=True),
        "left_align": Alignment(horizontal="left", vertical="center", wrap_text=True),
        "thin_border": Border(
            left=Side(style="thin"), right=Side(style="thin"),
            top=Side(style="thin"), bottom=Side(style="thin")
        ),
    }


def _set_header(ws, row, values, styles, fill_key="header_fill", start_col=1):
    for c, v in enumerate(values, start_col):
        cell = ws.cell(row=row, column=c, value=v)
        cell.font = styles["header_font"]
        cell.fill = styles[fill_key]
        cell.alignment = styles["center_align"]
        cell.border = styles["thin_border"]


def _write_row(ws, row, values, styles, align="left", is_input=False, is_note=False):
    align_obj = styles["left_align"] if align == "left" else styles["center_align"]
    for c, v in enumerate(values, 1):
        cell = ws.cell(row=row, column=c, value=v)
        cell.font = styles["blue_font"] if is_input else (styles["gray_font"] if is_note else styles["black_font"])
        cell.alignment = align_obj
        cell.border = styles["thin_border"]


def _adjust_columns(ws, min_width=10, max_width=60):
    for column in ws.columns:
        max_len = 0
        col_letter = get_column_letter(column[0].column)
        for cell in column:
            try:
                val = str(cell.value) if cell.value is not None else ""
                max_len = max(max_len, len(val))
            except Exception:
                pass
        width = min(max(max_len + 2, min_width), max_width)
        ws.column_dimensions[col_letter].width = width


def _freeze_header(ws, row=2):
    ws.freeze_panes = f"A{row}"


def _add_disclaimer(ws, start_row, styles, cols):
    note = (
        "【人工智能生成-需人工核验】本表内容仅供招聘管理参考，不构成正式劳动用工、薪酬决策或法律文件。"
        "最终招聘方案、录用决定及薪酬待遇须由企业 HR、用人部门及管理层审核确认。"
    )
    cell = ws.cell(row=start_row, column=1, value=note)
    cell.font = styles["orange_font"]
    cell.alignment = styles["left_align"]
    if cols > 1:
        ws.merge_cells(start_row=start_row, start_column=1, end_row=start_row, end_column=cols)


# ---------------------------------------------------------------------------
# Default data loaders
# ---------------------------------------------------------------------------
def _default_data():
    return {
        "city_multipliers": {
            "一线城市": {"multiplier": 1.3, "cities": ["北京", "上海", "深圳", "广州"]},
            "新一线城市": {"multiplier": 1.15, "cities": ["杭州", "成都", "南京", "武汉", "重庆", "天津", "苏州", "西安", "长沙", "青岛", "郑州", "宁波", "东莞", "佛山", "合肥"]},
            "二线城市": {"multiplier": 1.0, "cities": ["无锡", "厦门", "福州", "济南", "沈阳", "大连", "昆明", "哈尔滨", "长春", "石家庄", "南昌", "贵阳", "南宁", "海口", "兰州", "银川", "西宁", "乌鲁木齐", "呼和浩特", "太原"]},
            "三线城市": {"multiplier": 0.85, "cities": ["珠海", "惠州", "中山", "南通", "常州", "徐州", "扬州", "绍兴", "嘉兴", "台州", "金华", "温州", "泉州", "烟台", "威海", "潍坊", "淄博", "临沂", "唐山", "秦皇岛", "保定", "邯郸", "洛阳", "南阳", "襄阳", "宜昌", "岳阳", "常德", "株洲", "湘潭", "柳州", "桂林", "三亚", "拉萨", "大庆", "包头", "鄂尔多斯", "芜湖", "蚌埠", "滁州", "马鞍山", "安庆", "阜阳", "宿州", "六安", "淮安", "盐城", "连云港", "宿迁", "镇江", "泰州", "湖州", "衢州", "舟山", "丽水"]},
            "四线及以下": {"multiplier": 0.75, "cities": []},
        },
        "level_bandwidth": {
            "P1": {"name": "初级专员", "salary_min": 3500, "salary_max": 5000, "performance_months": 0.5},
            "P2": {"name": "专员", "salary_min": 5000, "salary_max": 8000, "performance_months": 0.8},
            "P3": {"name": "高级专员", "salary_min": 7000, "salary_max": 12000, "performance_months": 1.0},
            "P4": {"name": "主管", "salary_min": 10000, "salary_max": 16000, "performance_months": 1.5},
            "P5": {"name": "经理", "salary_min": 15000, "salary_max": 25000, "performance_months": 2.0},
            "P6": {"name": "高级经理", "salary_min": 22000, "salary_max": 35000, "performance_months": 2.5},
            "P7": {"name": "总监", "salary_min": 30000, "salary_max": 50000, "performance_months": 3.0},
            "P8": {"name": "高级总监", "salary_min": 45000, "salary_max": 70000, "performance_months": 4.0},
            "P9": {"name": "VP/高管", "salary_min": 60000, "salary_max": 100000, "performance_months": 6.0},
        },
        "position_kpi": {
            "电商运营助理": ["店铺销售额达成率", "流量转化率", "客单价维护", "活动执行完成率", "商品上架准确率", "客户好评率"],
            "电商运营": ["GMV 达成率", "ROI", "用户复购率", "流量成本管控", "活动策划与执行", "库存周转"],
            "产品经理": ["产品需求交付准时率", "用户满意度", "项目上线成功率", "需求变更率", "数据指标达成率"],
            "软件工程师": ["代码交付准时率", "Bug 逃逸率", "代码评审通过率", "技术债务清理率", "线上故障数"],
            "服装缝纫组组长": ["产量达成率", "质量合格率", "返修率", "人员出勤率", "员工流失率", "物料损耗率"],
            "通用": ["工作目标达成率", "工作质量", "团队协作", "执行力", "学习能力", "责任心"],
        },
        "competency": {
            "通用": ["执行力", "沟通能力", "学习能力", "团队协作", "责任心"],
            "运营类": ["数据分析", "活动策划", "用户洞察", "平台工具", "执行力"],
            "产品类": ["需求分析", "用户研究", "项目管理", "数据驱动", "跨部门协作"],
            "技术类": ["专业技术", "代码质量", "问题解决", "系统设计", "团队协作"],
            "生产类": ["生产执行", "质量管控", "人员管理", "成本意识", "安全意识"],
            "销售类": ["客户开发", "商务谈判", "结果导向", "抗压能力", "市场敏锐度"],
        },
        "channel_defaults": [
            {"channel": "招聘网站（BOSS/智联/前程）", "applications": 80, "screen_pass": 25, "first_interview": 12, "second_interview": 6, "offer": 3, "hire": 2},
            {"channel": "内部推荐", "applications": 20, "screen_pass": 12, "first_interview": 8, "second_interview": 5, "offer": 3, "hire": 2},
            {"channel": "猎头", "applications": 8, "screen_pass": 6, "first_interview": 5, "second_interview": 3, "offer": 2, "hire": 1},
            {"channel": "校园招聘", "applications": 30, "screen_pass": 10, "first_interview": 6, "second_interview": 3, "offer": 2, "hire": 2},
            {"channel": "官网/公众号", "applications": 15, "screen_pass": 4, "first_interview": 2, "second_interview": 1, "offer": 0, "hire": 0},
        ],
    }


def load_data():
    path = DATA_DIR / "recruitment_defaults.json"
    if path.exists():
        try:
            return load_json(path)
        except Exception:
            pass
    return _default_data()


def get_city_multiplier(city, data):
    for tier, info in data["city_multipliers"].items():
        if city in info["cities"]:
            return tier, info["multiplier"]
    return "default", 1.0


def get_level_info(level, data):
    return data["level_bandwidth"].get(level, data["level_bandwidth"]["P2"])


def detect_position_type(position):
    keywords = {
        "技术类": ["工程师", "开发", "架构", "算法", "测试", "运维", "DBA"],
        "产品类": ["产品经理", "产品专员", "产品助理", "产品总监"],
        "运营类": ["运营", "新媒体", "内容", "活动", "用户", "社群", "电商"],
        "销售类": ["销售", "商务", "BD", "客户经理", "渠道"],
        "生产类": ["生产", "制造", "缝纫", "裁剪", "质检", "车间", "班组长"],
    }
    for ptype, words in keywords.items():
        if any(w in position for w in words):
            return ptype
    return "通用"


def get_kpis(position, data):
    for key, kpis in data["position_kpi"].items():
        if key in position:
            return kpis
    return data["position_kpi"]["通用"]


def get_competencies(position, data):
    ptype = detect_position_type(position)
    return data["competency"].get(ptype, data["competency"]["通用"])


# ---------------------------------------------------------------------------
# Excel generators
# ---------------------------------------------------------------------------
def build_jd_sheet(ws, args, styles):
    ws.title = "岗位说明书（JD）"
    info = [
        ("岗位名称", args.position),
        ("所属部门", args.department),
        ("直接上级", args.reports_to),
        ("招聘人数", args.headcount),
        ("工作城市", args.city),
        ("职级", args.level),
        ("薪资范围", f"{args.salary_min} - {args.salary_max} 元/月"),
        ("紧急程度", args.urgency),
        ("编制日期", datetime.now().strftime("%Y-%m-%d")),
    ]
    ws.cell(row=1, column=1, value="岗位说明书（JD）").font = styles["title_font"]
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=2)
    ws.cell(row=1, column=1).alignment = styles["center_align"]

    for i, (k, v) in enumerate(info, 2):
        _write_row(ws, i, [k, v], styles, align="left", is_input=(k != "编制日期"))

    sections = [
        ("岗位目的", f"负责 {args.position} 相关工作，支撑 {args.department} 业务目标达成，提升团队运营效率与业务结果。"),
        ("核心职责", "1. 根据部门目标制定并执行岗位工作计划；\n2. 完成日常业务运营/技术支持/项目推进工作；\n3. 协同上下游团队，确保信息流转与任务交付质量；\n4. 持续优化工作流程，输出经验沉淀；\n5. 完成上级交办的其他工作任务。"),
        ("工作关系", f"内部：{args.department}、产品/技术/运营/销售等相关部门；外部：客户、供应商、合作伙伴、招聘渠道等。"),
        ("绩效目标", "依据公司绩效管理制度及岗位 KPI 指标执行，详见《结构化面试评估表》与绩效合同。"),
        ("备注", "本 JD 为模板化生成，具体职责与权限由用人部门根据实际业务调整。"),
    ]
    row = 12
    for title, content in sections:
        ws.cell(row=row, column=1, value=title).font = styles["section_font"]
        ws.cell(row=row, column=2, value=content).alignment = styles["left_align"]
        ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=4)
        row += 1

    _add_disclaimer(ws, row + 1, styles, 4)
    _adjust_columns(ws)


def build_qualification_sheet(ws, args, styles):
    ws.title = "任职资格与胜任力模型"
    ws.cell(row=1, column=1, value="任职资格与胜任力模型").font = styles["title_font"]
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=4)
    ws.cell(row=1, column=1).alignment = styles["center_align"]

    headers = ["维度", "要求项", "具体要求", "重要性"]
    _set_header(ws, 2, headers, styles)

    rows = [
        ("基本条件", "学历要求", "大专及以上学历，相关专业优先", "必备"),
        ("基本条件", "工作经验", f"{args.level} 职级对应经验要求，详见职级定义", "必备"),
        ("基本条件", "专业背景", "与岗位业务相关的专业或从业背景", "优先"),
        ("专业知识", "行业知识", "了解所在行业基本业务流程与市场动态", "必备"),
        ("专业知识", "岗位知识", "掌握岗位所需的专业工具、方法与规范", "必备"),
        ("专业技能", "工具技能", "熟练使用岗位相关办公软件与业务系统", "必备"),
        ("专业技能", "数据分析", "具备基础的数据收集、整理与分析能力", "优先"),
        ("通用能力", "沟通表达", "能够清晰表达观点，有效协同团队", "必备"),
        ("通用能力", "执行能力", "目标导向，能够按时保质完成任务", "必备"),
        ("通用能力", "学习能力", "主动学习新知识并快速应用于工作", "优先"),
        ("价值观", "责任心", "对工作结果负责，敢于承担", "必备"),
        ("价值观", "团队协作", "积极配合团队，共同达成目标", "必备"),
    ]
    for i, row_data in enumerate(rows, 3):
        _write_row(ws, i, row_data, styles)

    ws.cell(row=16, column=1, value="胜任力模型").font = styles["section_font"]
    ws.merge_cells(start_row=16, start_column=1, end_row=16, end_column=4)

    comp_headers = ["胜任力维度", "行为描述（优秀）", "行为描述（达标）", "行为描述（待提升）"]
    _set_header(ws, 17, comp_headers, styles)

    comps = get_competencies(args.position, load_data())
    comp_rows = [
        ("执行力", "主动推进复杂任务，超预期交付", "按计划完成目标，遇到障碍及时汇报", "需要督促，交付质量不稳定"),
        ("沟通能力", "跨团队高效协调，化解冲突", "表达清晰，信息传递准确", "表达含糊，易造成误解"),
        ("学习能力", "快速掌握新领域并输出方法论", "能主动学习并应用新知识", "学习意愿弱，成长缓慢"),
        ("团队协作", "主动补位，带动团队氛围", "配合团队完成共同目标", "协作意识不足"),
        ("责任心", "对结果高度负责，主动兜底", "承担分内责任，不推诿", "责任意识薄弱"),
    ]
    # Map generic competency descriptions to detected position type extras
    extra_comps = {
        "数据分析": ("数据分析", "独立完成复杂数据分析并驱动决策", "能进行基础数据统计与呈现", "缺乏数据敏感度"),
        "活动策划": ("活动策划", "独立策划中大型活动并达成目标", "能按模板执行活动方案", "活动策划能力弱"),
        "需求分析": ("需求分析", "深入洞察需求并输出高质量 PRD", "能收集整理需求并文档化", "需求理解偏差大"),
        "专业技术": ("专业技术", "独立解决复杂技术问题", "能完成常规开发/运维任务", "技术基础薄弱"),
        "生产执行": ("生产执行", "高效组织生产并达成产量质量双目标", "能按计划完成生产任务", "执行效率低"),
    }
    used = set()
    out_rows = []
    for c in comps:
        if c in extra_comps:
            out_rows.append(extra_comps[c])
            used.add(c)
        elif c not in used:
            for r in comp_rows:
                if r[0] == c:
                    out_rows.append(r)
                    used.add(c)
                    break
    # Fill remaining with generic if needed
    for r in comp_rows:
        if len(out_rows) >= 6:
            break
        if r[0] not in used:
            out_rows.append(r)
            used.add(r[0])

    for i, row_data in enumerate(out_rows, 18):
        _write_row(ws, i, row_data, styles)

    _add_disclaimer(ws, 26, styles, 4)
    _adjust_columns(ws)


def build_interview_sheet(ws, args, styles):
    ws.title = "结构化面试评估表"
    ws.cell(row=1, column=1, value=f"结构化面试评估表 - {args.position}").font = styles["title_font"]
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=8)
    ws.cell(row=1, column=1).alignment = styles["center_align"]

    info = [
        ("岗位名称", args.position),
        ("所属部门", args.department),
        ("职级", args.level),
        ("候选人姓名", ""),
        ("面试官", ""),
        ("面试日期", ""),
    ]
    for i, (k, v) in enumerate(info, 2):
        ws.cell(row=i, column=1, value=k).font = styles["black_font"]
        ws.cell(row=i, column=2, value=v).font = styles["blue_font"]
        ws.merge_cells(start_row=i, start_column=2, end_row=i, end_column=4)

    headers = ["序号", "评估维度", "权重", "评分标准", "面试官评分", "加权得分", "考察问题示例"]
    _set_header(ws, 9, headers, styles)

    kpis = get_kpis(args.position, load_data())
    # Ensure at least 6 dimensions including competencies
    comps = get_competencies(args.position, load_data())
    dimensions = []
    for k in kpis[:3]:
        dimensions.append((k, 15))
    for c in comps[:3]:
        dimensions.append((c, 10))
    # Add fixed general dimensions
    dimensions.append(("岗位匹配度", 15))
    dimensions.append(("稳定性与动机", 10))
    dimensions.append(("文化价值观", 10))
    # Normalize weights to sum 100
    total = sum(w for _, w in dimensions)
    if total != 100:
        dimensions = [(d, round(w * 100 / total)) for d, w in dimensions]
    # Adjust rounding errors
    diff = 100 - sum(w for _, w in dimensions)
    if dimensions:
        dimensions[-1] = (dimensions[-1][0], dimensions[-1][1] + diff)

    score_criteria = [
        "5=远超要求；4=符合要求；3=基本符合；2=部分符合；1=不符合",
        "5=卓越；4=良好；3=达标；2=待提升；1=不足",
        "5=深刻理解；4=熟悉；3=了解；2=薄弱；1=不了解",
    ]

    questions = [
        "请分享你过往在该领域最得意的案例及关键数据。",
        "描述一次你解决复杂问题的过程。",
        "如果要在 3 个月内达成某项目标，你会如何拆解并执行？",
        "请举例说明你与跨部门协作的经历。",
        "你对我们公司及这个岗位有哪些了解？",
        "你的职业规划是什么？为什么考虑这个机会？",
        "请举例说明你认同的团队文化或价值观。",
    ]

    for i, (dim, weight) in enumerate(dimensions, 10):
        score_row = i - 9
        criteria = score_criteria[(i - 10) % len(score_criteria)]
        question = questions[(i - 10) % len(questions)]
        ws.cell(row=i, column=1, value=score_row)
        ws.cell(row=i, column=2, value=dim)
        ws.cell(row=i, column=3, value=weight)
        ws.cell(row=i, column=4, value=criteria)
        ws.cell(row=i, column=5, value="").font = styles["blue_font"]
        ws.cell(row=i, column=6, value=f"=E{i}*C{i}/100")
        ws.cell(row=i, column=7, value=question)
        for c in range(1, 8):
            ws.cell(row=i, column=c).border = styles["thin_border"]
            ws.cell(row=i, column=c).alignment = styles["left_align"]

    total_row = 10 + len(dimensions)
    ws.cell(row=total_row, column=1, value="合计")
    ws.cell(row=total_row, column=1).font = styles["section_font"]
    ws.merge_cells(start_row=total_row, start_column=1, end_row=total_row, end_column=2)
    ws.cell(row=total_row, column=3, value=f"=SUM(C10:C{total_row-1})")
    ws.cell(row=total_row, column=6, value=f"=SUM(F10:F{total_row-1})")
    for c in [3, 6]:
        ws.cell(row=total_row, column=c).font = styles["section_font"]
        ws.cell(row=total_row, column=c).border = styles["thin_border"]
    ws.cell(row=total_row, column=1).alignment = styles["center_align"]
    ws.cell(row=total_row, column=1).border = styles["thin_border"]

    conclusion_row = total_row + 2
    ws.cell(row=conclusion_row, column=1, value="面试结论")
    ws.cell(row=conclusion_row, column=1).font = styles["section_font"]
    ws.merge_cells(start_row=conclusion_row, start_column=1, end_row=conclusion_row, end_column=2)
    ws.cell(row=conclusion_row, column=3, value="强烈推荐 / 推荐 / 待定 / 不推荐")
    ws.merge_cells(start_row=conclusion_row, start_column=3, end_row=conclusion_row, end_column=6)
    ws.cell(row=conclusion_row, column=3).font = styles["blue_font"]
    ws.cell(row=conclusion_row, column=3).alignment = styles["center_align"]

    note_row = conclusion_row + 2
    _add_disclaimer(ws, note_row, styles, 7)
    _freeze_header(ws, row=9)
    _adjust_columns(ws)


def build_funnel_sheet(ws, args, styles):
    ws.title = "招聘漏斗与周报"
    ws.cell(row=1, column=1, value=f"招聘漏斗与周报 - {args.position}").font = styles["title_font"]
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=10)
    ws.cell(row=1, column=1).alignment = styles["center_align"]

    ws.cell(row=2, column=1, value=f"周期：{datetime.now().strftime('%Y年第%W周')}   招聘负责人：__________   填报日期：__________")
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=10)
    ws.cell(row=2, column=1).font = styles["gray_font"]

    headers = ["招聘渠道", "投递数", "初筛通过", "初试人数", "复试人数", "Offer 数", "入职人数", "初筛通过率", "初试通过率", "Offer 接受率"]
    _set_header(ws, 3, headers, styles)

    data = load_data()
    channels = data["channel_defaults"]
    for i, ch in enumerate(channels, 4):
        row = [
            ch["channel"],
            ch["applications"],
            ch["screen_pass"],
            ch["first_interview"],
            ch["second_interview"],
            ch["offer"],
            ch["hire"],
            f"=C{i}/B{i}",
            f"=D{i}/C{i}",
            f"=G{i}/F{i}" if ch["offer"] else "-",
        ]
        _write_row(ws, i, row, styles)
        ws.cell(row=i, column=8).number_format = "0.0%"
        ws.cell(row=i, column=9).number_format = "0.0%"
        ws.cell(row=i, column=10).number_format = "0.0%"

    total_row = 4 + len(channels)
    totals = ["合计"]
    for col in range(2, 8):
        col_letter = get_column_letter(col)
        totals.append(f"=SUM({col_letter}4:{col_letter}{total_row-1})")
    totals.extend([f"=C{total_row}/B{total_row}", f"=D{total_row}/C{total_row}", f"=G{total_row}/F{total_row}"])
    _write_row(ws, total_row, totals, styles, align="center")
    for c in range(8, 11):
        ws.cell(row=total_row, column=c).number_format = "0.0%"

    summary_row = total_row + 2
    ws.cell(row=summary_row, column=1, value="本周进展摘要").font = styles["section_font"]
    ws.merge_cells(start_row=summary_row, start_column=1, end_row=summary_row, end_column=2)
    ws.cell(row=summary_row, column=3, value="（请填写本周关键进展、卡点与下周计划）")
    ws.merge_cells(start_row=summary_row, start_column=3, end_row=summary_row, end_column=10)
    ws.cell(row=summary_row, column=3).font = styles["blue_font"]

    note_row = summary_row + 2
    _add_disclaimer(ws, note_row, styles, 10)
    _freeze_header(ws, row=3)
    _adjust_columns(ws)


def build_profile_sheet(ws, args, styles):
    ws.title = "人才画像卡"
    ws.cell(row=1, column=1, value=f"人才画像卡 - {args.position}").font = styles["title_font"]
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=3)
    ws.cell(row=1, column=1).alignment = styles["center_align"]

    sections = [
        ("基础画像", [
            ("目标岗位", args.position),
            ("所属部门", args.department),
            ("职级", args.level),
            ("工作城市", args.city),
            ("薪资带宽", f"{args.salary_min} - {args.salary_max} 元/月"),
        ]),
        ("关键经历", [
            ("行业背景", "具备与目标岗位相关的行业从业经验"),
            ("项目经历", "有独立负责或深度参与的成功项目案例"),
            ("团队规模", "有与目标职级匹配的团队协作或管理经验"),
        ]),
        ("核心能力", [
            ("必备能力", "；".join(get_competencies(args.position, load_data()))),
            ("加分能力", "数据分析、跨部门沟通、抗压能力、创新思维"),
        ]),
        ("动机与文化", [
            ("求职动机", "认可公司业务方向，有明确的职业发展诉求"),
            ("文化匹配", "务实、协作、结果导向，认同公司价值观"),
        ]),
        ("风险信号", [
            ("需关注", "频繁跳槽、履历断层、薪资期望过高、岗位理解偏差"),
        ]),
        ("吸引点", [
            ("核心吸引", "业务成长空间、薪酬竞争力、团队氛围、学习发展机会"),
        ]),
    ]

    row = 2
    for section_title, items in sections:
        ws.cell(row=row, column=1, value=section_title).font = styles["section_font"]
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=3)
        row += 1
        _set_header(ws, row, ["项目", "描述", "备注"], styles)
        row += 1
        for k, v in items:
            _write_row(ws, row, [k, v, ""], styles)
            row += 1
        row += 1

    _add_disclaimer(ws, row, styles, 3)
    _adjust_columns(ws)


def build_offer_sheet(ws, args, styles):
    ws.title = "Offer 薪资建议"
    ws.cell(row=1, column=1, value=f"Offer 薪资建议 - {args.position}").font = styles["title_font"]
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=5)
    ws.cell(row=1, column=1).alignment = styles["center_align"]

    data = load_data()
    tier, multiplier = get_city_multiplier(args.city, data)
    level = get_level_info(args.level, data)

    base_min = args.salary_min or int(level["salary_min"] * multiplier)
    base_max = args.salary_max or int(level["salary_max"] * multiplier)
    target = int((base_min + base_max) / 2)

    ws.cell(row=2, column=1, value="城市等级").font = styles["black_font"]
    ws.cell(row=2, column=2, value=tier)
    ws.cell(row=2, column=3, value="城市系数").font = styles["black_font"]
    ws.cell(row=2, column=4, value=multiplier)
    ws.cell(row=2, column=4).number_format = "0.00"

    ws.cell(row=3, column=1, value="职级").font = styles["black_font"]
    ws.cell(row=3, column=2, value=f"{args.level} / {level['name']}")
    ws.cell(row=3, column=3, value="职级带宽（税前月）").font = styles["black_font"]
    ws.cell(row=3, column=4, value=f"{level['salary_min']} - {level['salary_max']} 元")

    ws.cell(row=4, column=1, value="建议月薪下限").font = styles["black_font"]
    ws.cell(row=4, column=2, value=base_min)
    ws.cell(row=4, column=3, value="建议月薪上限").font = styles["black_font"]
    ws.cell(row=4, column=4, value=base_max)
    ws.cell(row=4, column=5, value="目标月薪").font = styles["black_font"]
    ws.cell(row=4, column=5, value=target)

    headers = ["薪酬结构", "占比建议", "月薪金额（元）", "年金额（元）", "备注"]
    _set_header(ws, 6, headers, styles)

    rows = [
        ("固定工资（基本工资+岗位工资）", "70%-80%", f"=ROUND(E4*0.75,0)", f"=C7*12", "按月固定发放"),
        ("绩效工资", "10%-20%", f"=ROUND(E4*0.15,0)", f"=C8*12", "按考核周期发放"),
        ("补贴（餐补/交通/通讯）", "5%-10%", f"=ROUND(E4*0.05,0)", f"=C9*12", "按月发放"),
        ("年终奖", "1-{pm} 个月".format(pm=level["performance_months"]), "-", f"=E4*{level['performance_months']}", "与公司及个人绩效挂钩"),
        ("年度总包（预估）", "-", "-", f"=SUM(D7:D10)", "不含股权/期权/特殊激励"),
    ]
    for i, row_data in enumerate(rows, 7):
        _write_row(ws, i, row_data, styles)

    market_row = 13
    ws.cell(row=market_row, column=1, value="市场竞争力评估").font = styles["section_font"]
    ws.merge_cells(start_row=market_row, start_column=1, end_row=market_row, end_column=5)
    market_items = [
        ("市场分位", "建议定位 50-75 分位，核心岗位可定位 75 分位以上"),
        ("调薪空间", "入职 6-12 个月后可依据绩效进行首次调薪评估"),
        ("保留策略", "结合职业发展路径、项目激励、培训资源进行综合保留"),
        ("审批要求", "超出职级带宽上限的 Offer 须提交 HRD/CEO 特批"),
    ]
    for i, (k, v) in enumerate(market_items, market_row + 1):
        _write_row(ws, i, [k, v, "", "", ""], styles)
        ws.merge_cells(start_row=i, start_column=2, end_row=i, end_column=5)

    _add_disclaimer(ws, market_row + 6, styles, 5)
    _adjust_columns(ws)


def generate_excel(args, output_path):
    wb = Workbook()
    styles = _base_styles()

    # Remove default sheet
    wb.remove(wb.active)

    build_jd_sheet(wb.create_sheet(), args, styles)
    build_qualification_sheet(wb.create_sheet(), args, styles)
    build_interview_sheet(wb.create_sheet(), args, styles)
    build_funnel_sheet(wb.create_sheet(), args, styles)
    build_profile_sheet(wb.create_sheet(), args, styles)
    build_offer_sheet(wb.create_sheet(), args, styles)

    wb.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# Word generator
# ---------------------------------------------------------------------------
def _set_doc_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(docx.oxml.ns.qn("w:eastAsia"), "Microsoft YaHei")
    return p


def _add_doc_paragraph(doc, text, bold=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(docx.oxml.ns.qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)
    run.bold = bold
    if align:
        p.alignment = align
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p


def generate_word(args, output_path):
    if not HAVE_DOCX:
        print("[WARN] python-docx 未安装，跳过 Word 文档生成。请执行：pip install python-docx")
        return None

    doc = Document()
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)

    # Title
    title = doc.add_paragraph()
    run = title.add_run(f"招聘套件 - {args.position}")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(docx.oxml.ns.qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta_text = f"岗位：{args.position}  |  部门：{args.department}  |  城市：{args.city}  |  职级：{args.level}  |  生成日期：{datetime.now().strftime('%Y-%m-%d')}"
    run = meta.add_run(meta_text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(docx.oxml.ns.qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # JD
    _set_doc_heading(doc, "一、岗位说明书（JD）", level=1)
    _set_doc_heading(doc, "1.1 岗位基本信息", level=2)
    info_items = [
        f"岗位名称：{args.position}",
        f"所属部门：{args.department}",
        f"直接上级：{args.reports_to}",
        f"招聘人数：{args.headcount} 人",
        f"工作城市：{args.city}",
        f"职级：{args.level}",
        f"薪资范围：{args.salary_min} - {args.salary_max} 元/月",
        f"紧急程度：{args.urgency}",
    ]
    for item in info_items:
        _add_doc_paragraph(doc, item)

    _set_doc_heading(doc, "1.2 岗位目的", level=2)
    _add_doc_paragraph(doc, f"负责 {args.position} 相关工作，支撑 {args.department} 业务目标达成，提升团队运营效率与业务结果。")

    _set_doc_heading(doc, "1.3 核心职责", level=2)
    duties = [
        "根据部门目标制定并执行岗位工作计划；",
        "完成日常业务运营/技术支持/项目推进工作；",
        "协同上下游团队，确保信息流转与任务交付质量；",
        "持续优化工作流程，输出经验沉淀；",
        "完成上级交办的其他工作任务。",
    ]
    for duty in duties:
        _add_doc_paragraph(doc, duty)

    _set_doc_heading(doc, "1.4 工作关系", level=2)
    _add_doc_paragraph(doc, f"内部：{args.department}、产品/技术/运营/销售等相关部门；外部：客户、供应商、合作伙伴、招聘渠道等。")

    _set_doc_heading(doc, "1.5 绩效目标", level=2)
    _add_doc_paragraph(doc, "依据公司绩效管理制度及岗位 KPI 指标执行，详见 Excel 工作簿《结构化面试评估表》。")

    # Interview outline
    _set_doc_heading(doc, "二、结构化面试提纲", level=1)
    _set_doc_heading(doc, "2.1 评估维度与权重", level=2)
    _add_doc_paragraph(doc, "请结合 Excel 工作簿《结构化面试评估表》进行打分，以下为各维度考察要点与示例问题：")

    kpis = get_kpis(args.position, load_data())
    comps = get_competencies(args.position, load_data())
    dimensions = []
    for k in kpis[:3]:
        dimensions.append(k)
    for c in comps[:3]:
        dimensions.append(c)
    dimensions.extend(["岗位匹配度", "稳定性与动机", "文化价值观"])

    questions = [
        "请分享你过往在该领域最得意的案例及关键数据。",
        "描述一次你解决复杂问题的过程。",
        "如果要在 3 个月内达成某项目标，你会如何拆解并执行？",
        "请举例说明你与跨部门协作的经历。",
        "你对我们公司及这个岗位有哪些了解？",
        "你的职业规划是什么？为什么考虑这个机会？",
        "请举例说明你认同的团队文化或价值观。",
    ]

    for i, dim in enumerate(dimensions[:6], 1):
        _set_doc_heading(doc, f"2.{i} {dim}", level=2)
        _add_doc_paragraph(doc, f"考察要点：评估候选人在 {dim} 方面的经验深度、方法论与成果。")
        _add_doc_paragraph(doc, f"示例问题：{questions[(i-1) % len(questions)]}")

    _set_doc_heading(doc, "2.7 面试结论", level=2)
    _add_doc_paragraph(doc, "综合评分 = ______ 分    结论：□ 强烈推荐  □ 推荐  □ 待定  □ 不推荐")
    _add_doc_paragraph(doc, "面试官签字：__________    日期：__________")

    # Offer template
    _set_doc_heading(doc, "三、录用通知书（Offer）模板", level=1)
    _add_doc_paragraph(doc, f"尊敬的 __________ 先生/女士：")
    _add_doc_paragraph(doc, f"我们很高兴地通知您，您已通过我司 {args.position} 岗位面试，现向您发出正式录用邀请。")

    _set_doc_heading(doc, "3.1 职位信息", level=2)
    offer_items = [
        f"职位名称：{args.position}",
        f"所属部门：{args.department}",
        f"汇报对象：{args.reports_to}",
        f"工作地点：{args.city}",
        f"职级：{args.level}",
        f"入职时间：____ 年 __ 月 __ 日",
    ]
    for item in offer_items:
        _add_doc_paragraph(doc, item)

    _set_doc_heading(doc, "3.2 薪酬待遇", level=2)
    _add_doc_paragraph(doc, f"税前月薪：__________ 元（位于 {args.salary_min} - {args.salary_max} 元带宽范围内）")
    _add_doc_paragraph(doc, "薪酬结构：固定工资 + 绩效工资 + 补贴 + 年终奖，具体比例与发放规则依据公司薪酬制度执行。")
    _add_doc_paragraph(doc, "五险一金：按当地政策及公司规定缴纳。")

    _set_doc_heading(doc, "3.3 回复与生效", level=2)
    _add_doc_paragraph(doc, "请您在收到本通知书后 __ 个工作日内书面回复是否接受。逾期未回复视为自动放弃。")
    _add_doc_paragraph(doc, "本通知书自双方签署劳动合同之日起生效。")

    doc.add_paragraph()
    _add_doc_paragraph(doc, "公司（盖章）：__________    授权人签字：__________    日期：__________", align=WD_ALIGN_PARAGRAPH.RIGHT)

    # Disclaimer
    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    run = disclaimer.add_run(
        "【人工智能生成-需人工核验】本套件内容仅供招聘管理参考，不构成正式劳动用工、薪酬决策或法律文件。"
        "最终招聘方案、录用决定及薪酬待遇须由企业 HR、用人部门及管理层依据当地法律法规和公司政策审核确认。"
    )
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(docx.oxml.ns.qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0x66, 0x00)
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def parse_args():
    parser = argparse.ArgumentParser(
        description="prajna 招聘助手：一键生成招聘套件（Excel + Word）"
    )
    parser.add_argument("--position", required=True, help="岗位名称")
    parser.add_argument("--department", default="人力资源部", help="所属部门")
    parser.add_argument("--city", default="广州", help="工作城市")
    parser.add_argument("--level", default="P2", help="职级 P1-P9")
    parser.add_argument("--salary-min", type=int, default=None, help="月薪下限（元）")
    parser.add_argument("--salary-max", type=int, default=None, help="月薪上限（元）")
    parser.add_argument("--reports-to", default="部门经理", help="汇报对象")
    parser.add_argument("--headcount", type=int, default=1, help="招聘人数")
    parser.add_argument("--urgency", default="中", choices=["高", "中", "低"], help="紧急程度")
    parser.add_argument("--output", default=str(DEFAULT_SAMPLES_DIR), help="输出目录或文件路径")
    parser.add_argument("--format", default="all", choices=["excel", "word", "all"], help="输出格式")
    return parser.parse_args()


def resolve_salary(args):
    data = load_data()
    level = get_level_info(args.level, data)
    tier, multiplier = get_city_multiplier(args.city, data)

    if args.salary_min is None:
        args.salary_min = int(level["salary_min"] * multiplier)
    if args.salary_max is None:
        args.salary_max = int(level["salary_max"] * multiplier)

    if args.salary_min > args.salary_max:
        args.salary_min, args.salary_max = args.salary_max, args.salary_min
    return args


def main():
    args = parse_args()
    args = resolve_salary(args)

    today = datetime.now().strftime("%Y%m%d")
    safe_position = _safe_filename(args.position)
    safe_city = _safe_filename(args.city)
    base_name = f"prajna_招聘套件_{safe_position}_{safe_city}_{today}"

    output_dir = Path(args.output)
    if output_dir.is_dir() or str(output_dir).endswith(("samples", "output", "out")):
        output_dir.mkdir(parents=True, exist_ok=True)
        excel_path = output_dir / f"{base_name}.xlsx"
        word_path = output_dir / f"{base_name}.docx"
    else:
        output_dir.parent.mkdir(parents=True, exist_ok=True)
        excel_path = output_dir.with_suffix(".xlsx")
        word_path = output_dir.with_suffix(".docx")

    generated = []
    if args.format in ("excel", "all"):
        generate_excel(args, str(excel_path))
        generated.append(str(excel_path))
        print(f"[OK] Excel 招聘套件已生成：{excel_path}")

    if args.format in ("word", "all"):
        wp = generate_word(args, str(word_path))
        if wp:
            generated.append(wp)
            print(f"[OK] Word 招聘套件已生成：{wp}")

    print("\n生成文件：")
    for f in generated:
        print(f"  - {f}")


if __name__ == "__main__":
    main()
