#!/usr/bin/env python3
"""
prajna-salary-template 生成器 v2.0.0
为任意行业/岗位生成完整的薪资结构 Excel 模板（6+ 工作表）。
支持：城市薪酬基准、职级带宽、现有薪酬表自动补全、自然语言岗位、
      合规校验、绩效-奖金联动、批量生成、图表、Word/PDF、历史版本。
"""

import argparse
import csv
import json
import os
import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

from openpyxl import Workbook, load_workbook
from openpyxl.chart import PieChart, BarChart, Reference
from openpyxl.chart.label import DataLabelList
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# Optional dependencies with graceful degrade
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAVE_DOCX = True
except ImportError:
    HAVE_DOCX = False

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    HAVE_REPORTLAB = True
except ImportError:
    HAVE_REPORTLAB = False


# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_DIR = SCRIPT_DIR.parent
DATA_DIR = SKILL_DIR / "data"
HISTORY_DIR = Path.home() / ".prajna" / "salary_template_history"


# ---------------------------------------------------------------------------
# Base presets (v1.0)
# ---------------------------------------------------------------------------
PRESETS = {
    "互联网-电商运营助理": {
        "industry": "互联网/电商",
        "position": "电商运营助理",
        "base_salary": 3500,
        "position_salary": 1500,
        "performance_std": 2000,
        "full_attendance": 200,
        "meal_subsidy": 300,
        "transport_subsidy": 200,
        "communication_subsidy": 100,
        "social_base": 7000,
        "housing_fund_rate": 0.05,
        "overtime_pay": 500,
        "actual_attendance_days": 21,
        "scheduled_days": 21.75,
        "kpi_weights": [
            ("销售额达成率", 0.25, 85),
            ("订单处理效率", 0.15, 90),
            ("客户满意度", 0.15, 88),
            ("库存准确率", 0.10, 95),
            ("活动执行完成率", 0.10, 100),
            ("数据报表准确率", 0.10, 90),
            ("工作态度与协作", 0.10, 92),
            ("学习成长", 0.05, 85),
        ],
        "bonus": {
            "sales_commission": 800,
            "performance_multiplier": 0.8,
            "quarter_bonus_months": 0.5,
        },
    },
    "互联网-产品经理": {
        "industry": "互联网/科技",
        "position": "产品经理",
        "base_salary": 5000,
        "position_salary": 3000,
        "performance_std": 4000,
        "full_attendance": 300,
        "meal_subsidy": 400,
        "transport_subsidy": 300,
        "communication_subsidy": 200,
        "social_base": 12000,
        "housing_fund_rate": 0.12,
        "overtime_pay": 0,
        "actual_attendance_days": 21,
        "scheduled_days": 21.75,
        "kpi_weights": [
            ("产品需求交付及时率", 0.25, 90),
            ("项目上线质量", 0.20, 88),
            ("用户满意度/NPS", 0.15, 85),
            ("跨部门协作评分", 0.15, 90),
            ("数据分析与决策支持", 0.15, 88),
            ("创新提案落地", 0.10, 80),
        ],
        "bonus": {
            "sales_commission": 0,
            "performance_multiplier": 0.85,
            "quarter_bonus_months": 1.0,
        },
    },
    "互联网-软件工程师": {
        "industry": "互联网/科技",
        "position": "软件工程师",
        "base_salary": 6000,
        "position_salary": 4000,
        "performance_std": 3000,
        "full_attendance": 300,
        "meal_subsidy": 400,
        "transport_subsidy": 300,
        "communication_subsidy": 100,
        "social_base": 13000,
        "housing_fund_rate": 0.12,
        "overtime_pay": 0,
        "actual_attendance_days": 21,
        "scheduled_days": 21.75,
        "kpi_weights": [
            ("代码交付准时率", 0.25, 92),
            ("Bug 率与代码质量", 0.25, 88),
            ("技术文档完整性", 0.15, 85),
            ("系统稳定性保障", 0.15, 95),
            ("团队协作与知识分享", 0.10, 90),
            ("技术学习与成长", 0.10, 85),
        ],
        "bonus": {
            "sales_commission": 0,
            "performance_multiplier": 0.9,
            "quarter_bonus_months": 1.0,
        },
    },
    "制造业-生产主管": {
        "industry": "制造业",
        "position": "生产主管",
        "base_salary": 4000,
        "position_salary": 2500,
        "performance_std": 2500,
        "full_attendance": 300,
        "meal_subsidy": 350,
        "transport_subsidy": 250,
        "communication_subsidy": 150,
        "social_base": 9000,
        "housing_fund_rate": 0.08,
        "overtime_pay": 800,
        "actual_attendance_days": 22,
        "scheduled_days": 21.75,
        "kpi_weights": [
            ("生产计划达成率", 0.25, 90),
            ("产品合格率", 0.20, 95),
            ("安全生产事故", 0.20, 100),
            ("成本控制", 0.15, 88),
            ("团队人员流失率", 0.10, 85),
            ("5S 现场管理", 0.10, 90),
        ],
        "bonus": {
            "sales_commission": 0,
            "performance_multiplier": 0.88,
            "quarter_bonus_months": 0.8,
        },
    },
    "制造业-质检员": {
        "industry": "制造业",
        "position": "质检员",
        "base_salary": 2800,
        "position_salary": 1200,
        "performance_std": 1000,
        "full_attendance": 200,
        "meal_subsidy": 300,
        "transport_subsidy": 200,
        "communication_subsidy": 50,
        "social_base": 5000,
        "housing_fund_rate": 0.05,
        "overtime_pay": 600,
        "actual_attendance_days": 22,
        "scheduled_days": 21.75,
        "kpi_weights": [
            ("检验准确率", 0.30, 96),
            ("漏检率控制", 0.25, 95),
            ("检验效率", 0.20, 90),
            ("质量异常反馈及时率", 0.15, 92),
            ("工作纪律", 0.10, 90),
        ],
        "bonus": {
            "sales_commission": 0,
            "performance_multiplier": 0.85,
            "quarter_bonus_months": 0.5,
        },
    },
    "零售业-门店店长": {
        "industry": "零售/连锁",
        "position": "门店店长",
        "base_salary": 3500,
        "position_salary": 2500,
        "performance_std": 3000,
        "full_attendance": 300,
        "meal_subsidy": 300,
        "transport_subsidy": 200,
        "communication_subsidy": 200,
        "social_base": 9000,
        "housing_fund_rate": 0.08,
        "overtime_pay": 0,
        "actual_attendance_days": 21,
        "scheduled_days": 21.75,
        "kpi_weights": [
            ("门店销售额达成率", 0.30, 88),
            ("门店利润率", 0.20, 85),
            ("库存周转率", 0.15, 90),
            ("顾客满意度", 0.15, 92),
            ("团队管理与培训", 0.10, 88),
            ("损耗控制", 0.10, 90),
        ],
        "bonus": {
            "sales_commission": 1500,
            "performance_multiplier": 0.8,
            "quarter_bonus_months": 1.0,
        },
    },
    "零售业-销售顾问": {
        "industry": "零售/连锁",
        "position": "销售顾问",
        "base_salary": 2500,
        "position_salary": 1000,
        "performance_std": 1500,
        "full_attendance": 200,
        "meal_subsidy": 250,
        "transport_subsidy": 150,
        "communication_subsidy": 50,
        "social_base": 4500,
        "housing_fund_rate": 0.05,
        "overtime_pay": 0,
        "actual_attendance_days": 21,
        "scheduled_days": 21.75,
        "kpi_weights": [
            ("个人销售额达成率", 0.35, 90),
            ("客单价/连带率", 0.20, 88),
            ("会员开发与维护", 0.15, 85),
            ("顾客服务满意度", 0.15, 92),
            ("陈列与库存管理", 0.10, 90),
            ("学习与成长", 0.05, 85),
        ],
        "bonus": {
            "sales_commission": 1200,
            "performance_multiplier": 0.85,
            "quarter_bonus_months": 0.5,
        },
    },
    "金融业-客户经理": {
        "industry": "金融/银行",
        "position": "客户经理",
        "base_salary": 4000,
        "position_salary": 3000,
        "performance_std": 4000,
        "full_attendance": 300,
        "meal_subsidy": 400,
        "transport_subsidy": 300,
        "communication_subsidy": 300,
        "social_base": 12000,
        "housing_fund_rate": 0.12,
        "overtime_pay": 0,
        "actual_attendance_days": 21,
        "scheduled_days": 21.75,
        "kpi_weights": [
            ("新增客户数/资产规模", 0.30, 85),
            ("销售业绩达成率", 0.25, 88),
            ("客户满意度与投诉率", 0.20, 92),
            ("合规操作", 0.15, 100),
            ("产品知识考核", 0.10, 90),
        ],
        "bonus": {
            "sales_commission": 2500,
            "performance_multiplier": 0.8,
            "quarter_bonus_months": 1.5,
        },
    },
    "餐饮业-餐厅经理": {
        "industry": "餐饮/服务",
        "position": "餐厅经理",
        "base_salary": 3500,
        "position_salary": 2500,
        "performance_std": 2500,
        "full_attendance": 300,
        "meal_subsidy": 500,
        "transport_subsidy": 200,
        "communication_subsidy": 100,
        "social_base": 8000,
        "housing_fund_rate": 0.05,
        "overtime_pay": 0,
        "actual_attendance_days": 26,
        "scheduled_days": 26,
        "kpi_weights": [
            ("营业额达成率", 0.30, 90),
            ("毛利率控制", 0.20, 85),
            ("顾客满意度", 0.20, 92),
            ("食品安全与卫生", 0.15, 95),
            ("团队稳定与培训", 0.10, 88),
            ("损耗与成本控制", 0.05, 90),
        ],
        "bonus": {
            "sales_commission": 1000,
            "performance_multiplier": 0.85,
            "quarter_bonus_months": 1.0,
        },
    },
    "通用-行政专员": {
        "industry": "综合/行政",
        "position": "行政专员",
        "base_salary": 3000,
        "position_salary": 1500,
        "performance_std": 1500,
        "full_attendance": 200,
        "meal_subsidy": 300,
        "transport_subsidy": 200,
        "communication_subsidy": 100,
        "social_base": 6000,
        "housing_fund_rate": 0.05,
        "overtime_pay": 0,
        "actual_attendance_days": 21,
        "scheduled_days": 21.75,
        "kpi_weights": [
            ("行政事务完成及时率", 0.25, 92),
            ("办公成本控制", 0.20, 90),
            ("员工服务满意度", 0.20, 90),
            ("资产管理准确率", 0.15, 95),
            ("会议与活动组织", 0.10, 88),
            ("制度执行与合规", 0.10, 90),
        ],
        "bonus": {
            "sales_commission": 0,
            "performance_multiplier": 0.9,
            "quarter_bonus_months": 0.5,
        },
    },
}


DEFAULT_PARAMS = {
    "industry": "综合/行政",
    "position": "行政专员",
    "base_salary": 3000,
    "position_salary": 1500,
    "performance_std": 1500,
    "full_attendance": 200,
    "meal_subsidy": 300,
    "transport_subsidy": 200,
    "communication_subsidy": 100,
    "social_base": 6000,
    "housing_fund_rate": 0.05,
    "overtime_pay": 0,
    "actual_attendance_days": 21,
    "scheduled_days": 21.75,
    "kpi_weights": [
        ("工作完成质量", 0.25, 92),
        ("工作效率", 0.20, 88),
        ("团队协作", 0.15, 90),
        ("学习成长", 0.15, 85),
        ("出勤纪律", 0.15, 95),
        ("创新能力", 0.10, 80),
    ],
    "bonus": {
        "sales_commission": 0,
        "performance_multiplier": 0.9,
        "quarter_bonus_months": 0.5,
    },
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def load_json(path):
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def get_preset(name):
    return PRESETS.get(name, {}).copy() if name in PRESETS else None


def list_presets():
    return list(PRESETS.keys())


def _safe_filename(name):
    name = re.sub(r"[\\/:*?\"<>|\s]+", "_", name)
    return name.strip("_") or "template"


def _base_styles():
    return {
        "header_font": Font(name="Arial", bold=True, color="FFFFFF", size=11),
        "header_fill": PatternFill("solid", fgColor="1F4E78"),
        "blue_font": Font(name="Arial", color="0000FF", size=10),
        "black_font": Font(name="Arial", color="000000", size=10),
        "green_font": Font(name="Arial", color="008000", size=10),
        "red_font": Font(name="Arial", color="FF0000", size=10),
        "orange_font": Font(name="Arial", color="FF6600", size=10),
        "center_align": Alignment(horizontal="center", vertical="center", wrap_text=True),
        "left_align": Alignment(horizontal="left", vertical="center", wrap_text=True),
        "thin_border": Border(
            left=Side(style="thin"), right=Side(style="thin"),
            top=Side(style="thin"), bottom=Side(style="thin")
        ),
    }


def _set_header(ws, row, values, styles):
    for c, v in enumerate(values, 1):
        cell = ws.cell(row=row, column=c, value=v)
        cell.font = styles["header_font"]
        cell.fill = styles["header_fill"]
        cell.alignment = styles["center_align"]
        cell.border = styles["thin_border"]


def _write_row(ws, row, values, styles, is_input=False, is_link=False, is_warn=False):
    for c, v in enumerate(values, 1):
        cell = ws.cell(row=row, column=c, value=v)
        if is_input:
            cell.font = styles["blue_font"]
        elif is_warn:
            cell.font = styles["orange_font"]
        elif is_link:
            cell.font = styles["green_font"]
        else:
            cell.font = styles["black_font"]
        cell.alignment = styles["left_align"] if isinstance(v, str) else styles["center_align"]
        cell.border = styles["thin_border"]


# ---------------------------------------------------------------------------
# City / level / description resolution
# ---------------------------------------------------------------------------
def resolve_city(city, city_data):
    if not city:
        return None
    city = city.strip()
    detail = city_data.get("city_detail", {})
    if city in detail:
        return city, detail[city]
    for tier, info in city_data.get("tier", {}).items():
        if city in info.get("cities", []):
            return city, info
    return None


def get_city_multiplier(city_name, city_data):
    if not city_name:
        return 1.0
    for tier, info in city_data.get("tier", {}).items():
        if city_name in info.get("cities", []):
            return info.get("multiplier", 1.0)
    return 1.0


def apply_city_adjustment(params, city, city_data):
    if not city:
        return params
    city_name, city_info = resolve_city(city, city_data)
    if not city_name:
        return params

    multiplier = get_city_multiplier(city_name, city_data)
    params["city"] = city_name
    params["city_multiplier"] = multiplier

    for key in ["base_salary", "position_salary", "performance_std", "full_attendance",
                "meal_subsidy", "transport_subsidy", "communication_subsidy", "overtime_pay"]:
        if key in params and isinstance(params[key], (int, float)):
            params[key] = round(params[key] * multiplier)

    detail = city_data.get("city_detail", {}).get(city_name, {})
    if not detail:
        detail = city_data.get("city_detail", {}).get("default", {})
    if detail:
        gross = params.get("base_salary", 0) + params.get("position_salary", 0)
        proposed_base = max(gross, detail.get("social_base_min", gross))
        proposed_base = min(proposed_base, detail.get("social_base_max", proposed_base))
        params["social_base"] = proposed_base
        housing_min = detail.get("housing_base_min", gross)
        housing_max = detail.get("housing_base_max", gross)
        housing_base = max(min(proposed_base, housing_max), housing_min)
        params["housing_fund_base"] = housing_base
    return params


def apply_job_level(params, level, level_data):
    if not level:
        return params
    levels = level_data.get("levels", {})
    if level not in levels:
        return params
    info = levels[level]
    params["job_level"] = level
    params["job_level_name"] = info.get("name", level)

    params["base_salary"] = info.get("base_salary_mid", params.get("base_salary", 0))
    params["position_salary"] = info.get("position_salary_mid", params.get("position_salary", 0))
    params["performance_std"] = info.get("performance_std_mid", params.get("performance_std", 0))
    params["housing_fund_rate"] = info.get("housing_fund_rate", params.get("housing_fund_rate", 0.05))

    params["salary_bandwidth"] = {
        "base": (info.get("base_salary_min"), info.get("base_salary_mid"), info.get("base_salary_max")),
        "position": (info.get("position_salary_min"), info.get("position_salary_mid"), info.get("position_salary_max")),
        "performance": (info.get("performance_std_min"), info.get("performance_std_mid"), info.get("performance_std_max")),
    }
    params["bonus_bandwidth"] = {
        "quarter": (info.get("bonus_quarter_min"), info.get("bonus_quarter_mid"), info.get("bonus_quarter_max")),
        "annual": (info.get("bonus_annual_min"), info.get("bonus_annual_mid"), info.get("bonus_annual_max")),
    }
    return params


def infer_position(description, kpi_data, salary_data):
    """Rule-based inference from natural language description."""
    if not description:
        return None
    desc = description.lower()
    industries = kpi_data.get("industries", {})
    position_types = kpi_data.get("position_types", {})
    salary_templates = salary_data.get("position_types", {})

    matched_industry = "综合/行政"
    for industry, info in industries.items():
        for kw in info.get("keywords", []):
            if kw.lower() in desc:
                matched_industry = industry
                break
        if matched_industry != "综合/行政":
            break

    matched_type = "行政类"
    for ptype, keywords in position_types.items():
        for kw in keywords:
            if kw.lower() in desc:
                matched_type = ptype
                break
        if matched_type != "行政类":
            break

    industry_info = industries.get(matched_industry, {})
    kpi_templates = industry_info.get("kpi_templates", {})
    default_kpi = [["工作完成质量", 0.25, 90], ["工作效率", 0.20, 88], ["团队协作", 0.15, 90], ["学习成长", 0.15, 85], ["出勤纪律", 0.15, 95], ["创新能力", 0.10, 80]]
    kpi_weights = kpi_templates.get(matched_type, kpi_templates.get("通用", default_kpi))

    salary_template = salary_templates.get(matched_type, salary_templates.get("行政类", {}))
    position_name = description.split()[0] if description else "自定义岗位"

    return {
        "industry": matched_industry,
        "position": position_name,
        "kpi_weights": kpi_weights,
        "salary_template": salary_template,
    }


def apply_description_params(params, description, kpi_data, salary_data):
    inferred = infer_position(description, kpi_data, salary_data)
    if not inferred:
        return params
    params["industry"] = inferred["industry"]
    params["position"] = inferred["position"]
    params["kpi_weights"] = inferred["kpi_weights"]
    tmpl = inferred["salary_template"]
    for key, val in tmpl.items():
        params[key] = val
    params["bonus"] = {
        "sales_commission": tmpl.get("sales_commission", 0),
        "performance_multiplier": tmpl.get("performance_multiplier", 0.9),
        "quarter_bonus_months": tmpl.get("quarter_bonus_months", 0.5),
    }
    params.setdefault("actual_attendance_days", 21)
    params.setdefault("scheduled_days", 21.75)
    params["inferred_from_description"] = True
    return params


# ---------------------------------------------------------------------------
# Existing salary file parsing
# ---------------------------------------------------------------------------
SALARY_FIELD_ALIASES = {
    "base_salary": ["基本工资", "底薪", "基本薪资", "base", "基本工资标准"],
    "position_salary": ["岗位工资", "岗位薪资", "岗位津贴", "position", "岗位工资标准"],
    "performance_std": ["绩效工资", "绩效奖金", "绩效标准", "performance", "绩效工资标准"],
    "full_attendance": ["全勤奖", "全勤", "attendance"],
    "meal_subsidy": ["餐补", "餐饮补贴", "meal", "伙食补贴"],
    "transport_subsidy": ["交通补贴", "交通补助", "transport", "车补"],
    "communication_subsidy": ["通讯补贴", "通讯补助", "communication", "话费补贴"],
    "social_base": ["社保基数", "缴费基数", "social_base", "五险一金基数"],
    "housing_fund_rate": ["公积金比例", "公积金", "housing_rate"],
    "overtime_pay": ["加班费", "加班工资", "overtime"],
}


def parse_input_salary_file(input_path):
    """Parse an existing salary file and return recognized fields + diagnostics."""
    if not input_path or not os.path.exists(input_path):
        return {}, {}
    try:
        wb = load_workbook(input_path, data_only=True)
    except Exception as e:
        print(f"Warning: could not read input file {input_path}: {e}", file=sys.stderr)
        return {}, {}
    ws = wb.active
    recognized = {}
    diagnostics = {"recognized": [], "missing": list(SALARY_FIELD_ALIASES.keys())}

    for row in ws.iter_rows(min_row=1, max_row=20, min_col=1, max_col=10):
        for cell in row:
            if cell.value is None:
                continue
            text = str(cell.value).strip()
            for field, aliases in SALARY_FIELD_ALIASES.items():
                if any(alias in text for alias in aliases):
                    val_cell = ws.cell(row=cell.row, column=cell.column + 1)
                    if val_cell.value is not None:
                        try:
                            recognized[field] = float(val_cell.value)
                            if field in diagnostics["missing"]:
                                diagnostics["missing"].remove(field)
                            diagnostics["recognized"].append((field, val_cell.value))
                        except (ValueError, TypeError):
                            pass
    return recognized, diagnostics


# ---------------------------------------------------------------------------
# Parameter assembly
# ---------------------------------------------------------------------------
def build_params(args):
    city_data = load_json(DATA_DIR / "city_salary_multipliers.json")
    level_data = load_json(DATA_DIR / "job_level_bandwidth.json")
    kpi_data = load_json(DATA_DIR / "position_kpi_templates.json")
    salary_data = load_json(DATA_DIR / "position_salary_templates.json")

    recognized, diagnostics = {}, {}
    if getattr(args, "input", None):
        recognized, diagnostics = parse_input_salary_file(args.input)

    params = {}
    if getattr(args, "preset", None):
        preset = get_preset(args.preset)
        if not preset:
            print(f"Error: 未知预设 '{args.preset}'", file=sys.stderr)
            sys.exit(2)
        params = preset
    elif getattr(args, "description", None):
        params = apply_description_params(params, args.description, kpi_data, salary_data)
    elif getattr(args, "industry", None) or getattr(args, "position", None):
        synthetic = f"{getattr(args, 'industry', '') or ''}{getattr(args, 'position', '') or ''}"
        params = apply_description_params(params, synthetic, kpi_data, salary_data)
    else:
        params = {k: (v.copy() if isinstance(v, dict) else v) for k, v in DEFAULT_PARAMS.items()}

    for k, v in recognized.items():
        params[k] = v

    if getattr(args, "city", None):
        params = apply_city_adjustment(params, args.city, city_data)
    if getattr(args, "level", None):
        params = apply_job_level(params, args.level, level_data)

    if "bonus" not in params:
        params["bonus"] = {"sales_commission": 0, "performance_multiplier": 0.9, "quarter_bonus_months": 0.5}

    for k, v in DEFAULT_PARAMS.items():
        if k == "bonus":
            continue
        params.setdefault(k, v)

    return params, recognized, diagnostics


def build_params_from_fields(fields, city_data, level_data, kpi_data, salary_data):
    """Build params from a flat dict (used by batch mode)."""
    params = {}
    description = fields.get("description")
    if description:
        params = apply_description_params(params, description, kpi_data, salary_data)
    elif fields.get("industry") or fields.get("position"):
        synthetic = f"{fields.get('industry', '')}{fields.get('position', '')}"
        params = apply_description_params(params, synthetic, kpi_data, salary_data)
    else:
        params = {k: (v.copy() if isinstance(v, dict) else v) for k, v in DEFAULT_PARAMS.items()}

    if fields.get("city"):
        params = apply_city_adjustment(params, fields["city"], city_data)
    if fields.get("level"):
        params = apply_job_level(params, fields["level"], level_data)

    if "bonus" not in params:
        params["bonus"] = {"sales_commission": 0, "performance_multiplier": 0.9, "quarter_bonus_months": 0.5}
    for k, v in DEFAULT_PARAMS.items():
        if k == "bonus":
            continue
        params.setdefault(k, v)
    return params


# ---------------------------------------------------------------------------
# Excel building
# ---------------------------------------------------------------------------
def build_salary_workbook(params, output_path, output_format="excel"):
    """Main entry: build Excel and optionally Word/PDF."""
    wb = Workbook()
    styles = _base_styles()

    build_salary_structure(wb, params, styles)
    build_performance(wb, params, styles)
    build_adjustment(wb, params, styles)
    build_bonus(wb, params, styles)
    build_benefits(wb, params, styles)
    calc_info = _build_calculation_core(wb, params, styles)
    build_compliance(wb, params, styles, calc_info)
    build_score_coefficient(wb, params, styles)

    wb.save(output_path)

    generated = {"excel": output_path}
    if output_format in ("word", "all"):
        word_path = str(Path(output_path).with_suffix(".docx"))
        if HAVE_DOCX:
            generate_word(params, word_path)
            generated["word"] = word_path
        else:
            generated["word_warning"] = "python-docx not installed; Word output skipped"
            print(generated["word_warning"], file=sys.stderr)
    if output_format in ("pdf", "all"):
        pdf_path = str(Path(output_path).with_suffix(".pdf"))
        if HAVE_REPORTLAB:
            generate_pdf(params, pdf_path)
            generated["pdf"] = pdf_path
        else:
            generated["pdf_warning"] = "reportlab not installed; PDF output skipped"
            print(generated["pdf_warning"], file=sys.stderr)
    return generated


def build_salary_structure(wb, params, styles):
    ws = wb.active
    ws.title = "薪资结构"
    headers = ["序号", "薪资项目", "项目类型", "金额(元)", "薪酬带宽 Min", "薪酬带宽 Mid", "薪酬带宽 Max", "比例", "计算基数", "计算方式", "发放周期", "备注"]
    _set_header(ws, 1, headers, styles)

    base = params.get("base_salary", 0)
    pos = params.get("position_salary", 0)
    perf = params.get("performance_std", 0)
    full = params.get("full_attendance", 0)
    meal = params.get("meal_subsidy", 0)
    trans = params.get("transport_subsidy", 0)
    comm = params.get("communication_subsidy", 0)
    actual_days = params.get("actual_attendance_days", 21)
    scheduled_days = params.get("scheduled_days", 21.75)
    meal_actual = round(meal / scheduled_days * actual_days, 2)

    bandwidth = params.get("salary_bandwidth", {})
    b_base = bandwidth.get("base", ("-", "-", "-"))
    b_pos = bandwidth.get("position", ("-", "-", "-"))
    b_perf = bandwidth.get("performance", ("-", "-", "-"))

    rows = [
        [1, "基本工资", "固定薪资", base, b_base[0], b_base[1], b_base[2], "-", "合同约定", "按月固定发放", "月度", "当地最低工资以上"],
        [2, "岗位工资", "固定薪资", pos, b_pos[0], b_pos[1], b_pos[2], "-", "岗位职级", "按月固定发放", "月度", "根据岗位价值评估"],
        [3, "绩效工资", "浮动薪资", perf, b_perf[0], b_perf[1], b_perf[2], "-", "绩效考核得分", "绩效得分/100*标准", "月度", "与KPI挂钩"],
        [4, "加班工资", "浮动薪资", params.get("overtime_pay", 0), "-", "-", "-", "-", "实际加班时长", "按劳动法1.5/2/3倍", "月度", "以基本工资为基数"],
        [5, "全勤奖", "津贴补贴", full, "-", "-", "-", "-", "当月无缺勤", "满足条件全额发放", "月度", "迟到早退超3次取消"],
        [6, "餐补", "津贴补贴", meal_actual, "-", "-", "-", "-", "出勤天数", f"{meal}/{scheduled_days}*{actual_days}", "月度", "按实际出勤计算"],
        [7, "交通补贴", "津贴补贴", trans, "-", "-", "-", "-", "固定", "按月固定发放", "月度", "外勤较多可上浮"],
        [8, "通讯补贴", "津贴补贴", comm, "-", "-", "-", "-", "固定", "按月固定发放", "月度", "运营岗位标配"],
        [9, "五险一金(个人)", "代扣项", 0, "-", "-", "-", "-", "缴费基数*比例", "养老8%+医疗2%+失业0.5%+公积金", "月度", "单位另承担约30%"],
        [10, "个人所得税", "代扣项", 0, "-", "-", "-", "-", "应纳税所得额", "累计预扣法", "月度", "按个税法计算"],
        [11, "应发工资", "汇总", 0, "-", "-", "-", "-", "固定+浮动+补贴", "SUM(固定+浮动+补贴)", "月度", "税前总薪资"],
        [12, "实发工资", "汇总", 0, "-", "-", "-", "-", "应发-代扣项", "应发-五险一金-个税", "月度", "到手金额"],
    ]
    for i, row in enumerate(rows, 2):
        _write_row(ws, i, row, styles, is_input=(row[2] in ["固定薪资", "津贴补贴"] and row[0] <= 8))

    ws["D10"] = "='计算示例'!B20"
    ws["D11"] = "='计算示例'!B23"
    ws["D12"] = "='计算示例'!B19"
    ws["D13"] = "='计算示例'!B24"

    for col in range(1, len(headers) + 1):
        ws.column_dimensions[get_column_letter(col)].width = 14
    ws.column_dimensions["B"].width = 14
    ws.column_dimensions["L"].width = 22


def build_performance(wb, params, styles):
    ws = wb.create_sheet("绩效考核")
    headers = ["序号", "考核维度", "权重", "指标定义", "目标值", "评分标准", "数据来源", "考核周期", "备注"]
    _set_header(ws, 1, headers, styles)

    kpi_weights = params.get("kpi_weights", [["工作完成质量", 0.25, 90]])
    for idx, (name, weight, score) in enumerate(kpi_weights, 1):
        target = ">=90分" if weight >= 0.2 else ">=85分"
        standard = "达标100分；每低1分扣2-5分"
        source = "业务系统/主管评分"
        cycle = "月度"
        note = "核心指标" if weight >= 0.2 else "辅助指标"
        row = [idx, name, f"{weight:.0%}", name.replace("率", "实际/目标"), target, standard, source, cycle, note]
        _write_row(ws, idx + 1, row, styles, is_input=(idx == 1))

    for col in range(1, len(headers) + 1):
        ws.column_dimensions[get_column_letter(col)].width = 16
    ws.column_dimensions["D"].width = 26
    ws.column_dimensions["F"].width = 30
    ws.column_dimensions["I"].width = 20


def build_adjustment(wb, params, styles):
    ws = wb.create_sheet("调薪机制")
    headers = ["序号", "调薪类型", "适用条件", "调薪幅度", "生效时间", "审批流程", "备注"]
    _set_header(ws, 1, headers, styles)

    data = [
        [1, "试用期转正调薪", "试用期考核通过", "0-10%", "转正次月", "部门主管→HR→总经理", "根据试用期表现定档"],
        [2, "年度普调", "入职满1年且年度绩效≥B", "3-8%", "次年1月", "HR提案→管理层审批", "参考市场薪酬涨幅"],
        [3, "晋升调薪", "职级晋升", "10-25%", "晋升次月", "部门提名→晋升评审→总经理", "每晋升1级至少10%"],
        [4, "绩效调薪", "连续2个季度绩效A或年度绩效A", "5-15%", "绩效评定次月", "主管申请→HR审核→总经理", "不与晋升调薪重复享受"],
        [5, "市场薪酬调整", "岗位市场薪酬显著变化", "5-10%", "调整方案确认次月", "HR调研→管理层审批", "需有薪酬调研数据支撑"],
        [6, "特殊贡献调薪", "重大项目突破/特殊贡献", "10-30%", "贡献确认次月", "部门申请→总经理特批", "一事一议"],
        [7, "留任调薪", "核心人才面临外部高薪挖角", "10-20%", "审批通过后次月", "HR评估→总经理审批", "需签订服务期协议"],
    ]
    for i, row in enumerate(data, 2):
        _write_row(ws, i, row, styles)

    for col in range(1, len(headers) + 1):
        ws.column_dimensions[get_column_letter(col)].width = 18
    ws.column_dimensions["C"].width = 28
    ws.column_dimensions["G"].width = 26


def build_bonus(wb, params, styles):
    ws = wb.create_sheet("奖金方案")
    headers = ["序号", "奖金类型", "适用岗位", "计算基数", "计提比例/金额", "发放周期", "发放条件", "计算公式", "备注"]
    _set_header(ws, 1, headers, styles)

    position = params.get("position", "岗位")
    full = params.get("full_attendance", 200)
    bonus = params.get("bonus", {})

    bb = params.get("bonus_bandwidth", {})
    q_min, q_mid, q_max = bb.get("quarter", (0.5, 0.8, 1.5))
    a_min, a_mid, a_max = bb.get("annual", (1.0, 1.5, 3.0))

    rows = [
        [1, "月度销售提成", position, "个人负责销售额", "0.3%-0.8%", "月度", "销售额达标且回款完成", "销售额*提成比例", "阶梯提成"],
        [2, "月度绩效奖金", position, "绩效标准金额", "0-120%", "月度", "月度绩效得分≥60", "绩效标准*绩效系数", "与绩效考核表挂钩"],
        [3, "季度奖金", position, "月基本工资", f"{q_min}-{q_max}个月", "季度", "季度目标达成", f"基本工资*{q_mid}", "每季度评估"],
        [4, "年度奖金", position, "月基本工资", f"{a_min}-{a_max}个月", "年度", "年度绩效≥B且在职", f"基本工资*{a_mid}", "春节前发放"],
        [5, "项目奖金", position, "项目利润", "5%-15%", "项目结束", "项目验收通过", "项目利润*分配比例", "专项激励"],
        [6, "全勤奖", "全员", "固定金额", f"{full}元", "月度", "当月无缺勤迟到早退", f"固定{full}元", "超3次取消"],
        [7, "伯乐奖", "全员", "固定金额", "500-2000元", "实时", "推荐人才成功转正", "按岗位级别定", "转正后发放"],
        [8, "年终分红", "核心员工", "公司年度利润", "0.5%-2%", "年度", "入职满1年且绩效≥B", "利润池*个人分配比例", "仅限资深员工"],
    ]
    for i, row in enumerate(rows, 2):
        _write_row(ws, i, row, styles)

    for col in range(1, len(headers) + 1):
        ws.column_dimensions[get_column_letter(col)].width = 16
    ws.column_dimensions["C"].width = 18
    ws.column_dimensions["H"].width = 26
    ws.column_dimensions["I"].width = 22


def build_benefits(wb, params, styles):
    ws = wb.create_sheet("福利明细")
    headers = ["序号", "福利项目", "福利类型", "标准/金额", "享受条件", "发放方式", "备注"]
    _set_header(ws, 1, headers, styles)

    trans = params.get("transport_subsidy", 200)
    meal = params.get("meal_subsidy", 300)
    comm = params.get("communication_subsidy", 100)

    rows = [
        [1, "养老保险", "法定福利", "基数*16%(单位)+8%(个人)", "签订劳动合同", "社保代缴", "按当地社保政策"],
        [2, "医疗保险", "法定福利", "基数*8%(单位)+2%(个人)", "签订劳动合同", "社保代缴", "含生育保险"],
        [3, "失业保险", "法定福利", "基数*0.5%(单位)+0.5%(个人)", "签订劳动合同", "社保代缴", "按当地政策"],
        [4, "工伤保险", "法定福利", "基数*0.2%-1.4%(单位)", "签订劳动合同", "社保代缴", "按行业风险等级"],
        [5, "住房公积金", "法定福利", "基数*5%-12%(双边)", "签订劳动合同", "公积金代缴", "单位和个人同比例"],
        [6, "补充医疗保险", "补充福利", "约200元/人/年", "转正后", "商业保险", "覆盖门诊住院"],
        [7, "带薪年假", "法定福利", "5-15天/年", "入职满1年", "申请休假", "按工龄递增"],
        [8, "法定节假日", "法定福利", "11天/年", "全员", "带薪休假", "按国家规定"],
        [9, "节日福利", "补充福利", "300-500元/节", "全员", "实物/购物卡/现金", "春节/中秋/端午"],
        [10, "生日福利", "补充福利", "200元礼品/礼金", "全员", "生日当月发放", "或生日蛋糕券"],
        [11, "培训补贴", "补充福利", "1000-3000元/年", "年度培训计划内", "报销/直付", "外部课程/考证"],
        [12, "团建经费", "补充福利", "100-200元/人/月", "全员", "部门统筹", "季度或月度活动"],
        [13, "下午茶", "补充福利", "约20元/人/周", "全员", "行政采购", "每周固定"],
        [14, "年度体检", "补充福利", "500-1000元/人", "入职满1年", "合作医院", "每年1次"],
        [15, "交通补助", "津贴补贴", f"{trans}元/月", "外勤较多岗位", "随工资发放", "需提交外出记录"],
        [16, "餐补", "津贴补贴", f"{meal}元/月", "全员", "随工资发放", "按出勤折算"],
        [17, "通讯补贴", "津贴补贴", f"{comm}元/月", "运营/销售岗位", "随工资发放", "需报销话费发票"],
    ]
    for i, row in enumerate(rows, 2):
        _write_row(ws, i, row, styles)

    for col in range(1, len(headers) + 1):
        ws.column_dimensions[get_column_letter(col)].width = 18
    ws.column_dimensions["B"].width = 16
    ws.column_dimensions["D"].width = 26
    ws.column_dimensions["G"].width = 22


def build_score_coefficient(wb, params, styles):
    ws = wb.create_sheet("绩效系数对照表")
    headers = ["绩效得分区间", "绩效系数", "说明"]
    _set_header(ws, 1, headers, styles)

    rows = [
        ["< 60", "0.6", "不合格，绩效奖金打折"],
        ["60 - 70", "0.7", "待改进"],
        ["70 - 80", "0.8", "合格"],
        ["80 - 90", "0.9", "良好"],
        ["90 - 100", "1.0", "达标"],
        ["100 - 110", "1.1", "优秀"],
        [">= 110", "1.2", "卓越"],
    ]
    for i, row in enumerate(rows, 2):
        _write_row(ws, i, row, styles)

    for col in range(1, len(headers) + 1):
        ws.column_dimensions[get_column_letter(col)].width = 18
    ws.column_dimensions["C"].width = 24


def _build_calculation_core(wb, params, styles):
    ws = wb.create_sheet("计算示例")

    industry = params.get("industry", "通用行业")
    position = params.get("position", "通用岗位")
    city = params.get("city", "")
    city_multiplier = params.get("city_multiplier", 1.0)
    job_level = params.get("job_level", "")
    job_level_name = params.get("job_level_name", "")

    ws["A1"] = f"{industry} · {position} · 薪资计算示例"
    ws["A1"].font = Font(name="Arial", bold=True, size=14)
    ws.merge_cells("A1:F1")

    info = [
        ["员工姓名", "张三", "部门", f"{industry}事业部", "岗位", position],
        ["入职日期", "2026-03-01", "工龄", "1年", "职级", f"{job_level} {job_level_name}" if job_level else "P1"],
        ["计薪月份", datetime.now().strftime("%Y-%m"), "应出勤天数", params.get("scheduled_days", 21.75), "实际出勤天数", params.get("actual_attendance_days", 21)],
        ["基本工资", params.get("base_salary", 0), "岗位工资", params.get("position_salary", 0), "绩效工资标准", params.get("performance_std", 0)],
        ["全勤奖标准", params.get("full_attendance", 0), "餐补标准", params.get("meal_subsidy", 0), "交通补贴标准", params.get("transport_subsidy", 0)],
        ["通讯补贴标准", params.get("communication_subsidy", 0), "五险一金基数", params.get("social_base", 0), "公积金比例", params.get("housing_fund_rate", 0.05)],
    ]
    if city:
        info.append(["城市", city, "城市系数", city_multiplier, "", ""])

    for i, row in enumerate(info, 3):
        for j, v in enumerate(row, 1):
            cell = ws.cell(row=i, column=j, value=v)
            cell.font = styles["blue_font"] if j % 2 == 0 else styles["black_font"]
            cell.border = styles["thin_border"]
            cell.alignment = styles["center_align"]
            if j == 6 and i == 8:
                cell.number_format = "0%"

    ws["A10"] = "薪资项目"
    ws["B10"] = "金额(元)"
    ws["C10"] = "计算说明"
    for c in range(1, 4):
        cell = ws.cell(row=10, column=c)
        cell.font = styles["header_font"]
        cell.fill = styles["header_fill"]
        cell.alignment = styles["center_align"]
        cell.border = styles["thin_border"]

    kpi_weights = params.get("kpi_weights", [["工作完成质量", 0.25, 90]])
    perf_score_row = 11 + len(kpi_weights) + 1
    perf_score_cell = f"F{perf_score_row}"
    perf_std = params.get("performance_std", 0)
    meal_actual = round(params.get("meal_subsidy", 0) / params.get("scheduled_days", 21.75) * params.get("actual_attendance_days", 21), 2)

    rows = [
        ["基本工资", params.get("base_salary", 0), "固定"],
        ["岗位工资", params.get("position_salary", 0), "固定"],
        ["绩效工资", f"=MAX(0.6,MIN(1.2,{perf_score_cell}/100))*D8", f"由加权绩效得分驱动"],
        ["加班工资", params.get("overtime_pay", 0), "实际加班/值班"],
        ["全勤奖", params.get("full_attendance", 0), "当月无缺勤"],
        ["餐补", meal_actual, f"标准/应出勤*实际出勤"],
        ["交通补贴", params.get("transport_subsidy", 0), "固定"],
        ["通讯补贴", params.get("communication_subsidy", 0), "固定"],
        ["应发工资", "=SUM(B11:B18)", ""],
        ["五险一金(个人)", "=D8*(8%+2%+0.5%)+D8*F8", "养老+医疗+失业+公积金"],
        ["个税起征点", 5000, ""],
        ["应纳税所得额", "=MAX(B19-B20-B21,0)", ""],
        ["个人所得税", "=IF(B22<=0,0,IF(B22<=3000,B22*3%,IF(B22<=12000,B22*10%-210,IF(B22<=25000,B22*20%-1410,B22*25%-2660))))", "累计预扣法（简化）"],
        ["实发工资", "=B19-B20-B23", ""],
    ]
    for i, row in enumerate(rows, 11):
        for j, v in enumerate(row, 1):
            cell = ws.cell(row=i, column=j, value=v)
            if j == 1:
                cell.font = styles["black_font"]
            elif j == 2:
                if isinstance(v, str) and v.startswith("="):
                    cell.font = styles["green_font"]
                elif i in [11, 12, 13, 14, 15, 16, 17, 18]:
                    cell.font = styles["blue_font"]
                else:
                    cell.font = styles["black_font"]
            else:
                cell.font = styles["black_font"]
            cell.border = styles["thin_border"]
            cell.alignment = styles["left_align"] if j == 3 else styles["center_align"]

    ws["D10"] = "绩效考核示例"
    ws.merge_cells("D10:F10")
    ws["D10"].font = styles["header_font"]
    ws["D10"].fill = styles["header_fill"]
    ws["D10"].alignment = styles["center_align"]
    perf_headers = ["考核维度", "权重", "得分"]
    for j, h in enumerate(perf_headers, 4):
        cell = ws.cell(row=11, column=j, value=h)
        cell.font = styles["header_font"]
        cell.fill = styles["header_fill"]
        cell.alignment = styles["center_align"]
        cell.border = styles["thin_border"]

    perf_data_rows = []
    for name, weight, score in kpi_weights:
        perf_data_rows.append([name, weight, score])
    weighted_formula = f"=SUMPRODUCT(E12:E{11 + len(kpi_weights)},F12:F{11 + len(kpi_weights)})"
    perf_data_rows.append(["加权绩效得分", weighted_formula, ""])

    for i, row in enumerate(perf_data_rows, 12):
        for j, v in enumerate(row, 4):
            cell = ws.cell(row=i, column=j, value=v)
            cell.font = styles["green_font"] if isinstance(v, str) and v.startswith("=") else styles["black_font"]
            cell.border = styles["thin_border"]
            cell.alignment = styles["center_align"]
            if j == 5:
                cell.number_format = "0%"

    bonus_start = 26
    ws.cell(row=bonus_start, column=1, value="奖金计算示例")
    ws.merge_cells(start_row=bonus_start, start_column=1, end_row=bonus_start, end_column=3)
    ws.cell(row=bonus_start, column=1).font = styles["header_font"]
    ws.cell(row=bonus_start, column=1).fill = styles["header_fill"]
    ws.cell(row=bonus_start, column=1).alignment = styles["center_align"]

    bonus_headers = ["奖金项目", "金额(元)", "计算说明"]
    for j, h in enumerate(bonus_headers, 1):
        cell = ws.cell(row=bonus_start + 1, column=j, value=h)
        cell.font = styles["header_font"]
        cell.fill = styles["header_fill"]
        cell.alignment = styles["center_align"]
        cell.border = styles["thin_border"]

    quarter_bonus = round(params.get("base_salary", 0) * params.get("bonus", {}).get("quarter_bonus_months", 0.5), 2)
    sales_commission = params.get("bonus", {}).get("sales_commission", 0)
    full_attendance = params.get("full_attendance", 0)

    bonus_rows = [
        ["月度销售提成", sales_commission, "按实际业绩计提"],
        ["月度绩效奖金", "=B13", "绩效工资联动"],
        ["全勤奖", full_attendance, "当月无缺勤"],
        ["季度奖金(预估)", quarter_bonus, "基本工资*系数"],
        ["奖金合计", f"=SUM(B{bonus_start+2}:B{bonus_start+5})", ""],
    ]
    for i, row in enumerate(bonus_rows, bonus_start + 2):
        for j, v in enumerate(row, 1):
            cell = ws.cell(row=i, column=j, value=v)
            cell.font = styles["green_font"] if isinstance(v, str) and v.startswith("=") else styles["black_font"]
            cell.border = styles["thin_border"]
            cell.alignment = styles["center_align"]

    summary_start = bonus_start + 7
    ws.cell(row=summary_start, column=5, value="全年薪酬总包预估")
    ws.merge_cells(start_row=summary_start, start_column=5, end_row=summary_start, end_column=6)
    ws.cell(row=summary_start, column=5).font = styles["header_font"]
    ws.cell(row=summary_start, column=5).fill = styles["header_fill"]
    ws.cell(row=summary_start, column=5).alignment = styles["center_align"]

    summary_headers = ["项目", "金额(元/年)"]
    for j, h in enumerate(summary_headers, 5):
        cell = ws.cell(row=summary_start + 1, column=j, value=h)
        cell.font = styles["header_font"]
        cell.fill = styles["header_fill"]
        cell.alignment = styles["center_align"]
        cell.border = styles["thin_border"]

    summary_data = [
        ["年固定薪资", "=B19*12", ""],
        ["年浮动奖金(提成+绩效)", f"=(B{bonus_start+2}+B{bonus_start+3})*12", ""],
        ["年终奖(1个月)", "=B11+B12", ""],
        ["年福利(餐补+交通+通讯+节日)", f"=(B16+B17+B18+B{bonus_start+4})*12+1200", ""],
        ["全年总现金收入", f"=F{summary_start+2}+F{summary_start+3}+F{summary_start+4}+F{summary_start+5}", ""],
    ]
    for i, row in enumerate(summary_data, summary_start + 2):
        for j, v in enumerate(row, 5):
            cell = ws.cell(row=i, column=j, value=v)
            cell.font = styles["green_font"] if isinstance(v, str) and v.startswith("=") else styles["black_font"]
            cell.border = styles["thin_border"]
            cell.alignment = styles["center_align"]

    _add_charts(ws, bonus_start, summary_start)

    ws.column_dimensions["A"].width = 20
    ws.column_dimensions["B"].width = 16
    ws.column_dimensions["C"].width = 30
    ws.column_dimensions["D"].width = 18
    ws.column_dimensions["E"].width = 14
    ws.column_dimensions["F"].width = 18

    return {
        "gross_salary_cell": "B19",
        "social_insurance_cell": "B20",
        "tax_cell": "B23",
        "net_salary_cell": "B24",
        "taxable_income_cell": "B22",
        "perf_score_cell": perf_score_cell,
        "base_salary_cell": "B11",
        "position_salary_cell": "B12",
        "perf_std_cell": "D8",
        "housing_rate_cell": "F8",
        "bonus_start": bonus_start,
        "summary_start": summary_start,
    }


def _add_charts(ws, bonus_start, summary_start):
    pie = PieChart()
    pie.title = "薪资结构占比"
    labels = Reference(ws, min_col=1, min_row=11, max_row=18)
    data = Reference(ws, min_col=2, min_row=10, max_row=18)
    pie.add_data(data, titles_from_data=True)
    pie.set_categories(labels)
    pie.dataLabels = DataLabelList()
    pie.dataLabels.showPercent = True
    ws.add_chart(pie, "H10")

    bar = BarChart()
    bar.type = "col"
    bar.title = "固定薪酬 vs 浮动薪酬"
    bar.y_axis.title = "金额(元)"
    cats = Reference(ws, min_col=5, min_row=summary_start+2, max_row=summary_start+4)
    data = Reference(ws, min_col=6, min_row=summary_start+1, max_row=summary_start+4)
    bar.add_data(data, titles_from_data=True)
    bar.set_categories(cats)
    ws.add_chart(bar, "H25")


def build_compliance(wb, params, styles, calc_info):
    ws = wb.create_sheet("合规校验")
    headers = ["序号", "校验项", "规则", "当前值", "阈值", "结果", "说明"]
    _set_header(ws, 1, headers, styles)

    city = params.get("city", "")
    city_data = load_json(DATA_DIR / "city_salary_multipliers.json")
    city_detail = city_data.get("city_detail", {}).get(city, city_data.get("city_detail", {}).get("default", {}))
    min_wage = city_detail.get("min_wage", 2000)
    social_min = city_detail.get("social_base_min", 4000)
    social_max = city_detail.get("social_base_max", 20000)

    base_salary = params.get("base_salary", 0)
    social_base = params.get("social_base", 0)
    housing_rate = params.get("housing_fund_rate", 0.05)
    perf_std = params.get("performance_std", 0)
    gross_fixed = base_salary + params.get("position_salary", 0)

    rows = [
        [1, "最低工资合规", f"基本工资 ≥ {min_wage}元", base_salary, min_wage, f'=IF(D2>=E2,"PASS","BLOCK")', "低于最低工资为严重违规"],
        [2, "社保基数下限", f"社保基数 ≥ {social_min}元", social_base, social_min, f'=IF(D3>=E3,"PASS","WARN")', "低于下限需按最低基数缴纳"],
        [3, "社保基数上限", f"社保基数 ≤ {social_max}元", social_base, social_max, f'=IF(D4<=E4,"PASS","WARN")', "超过上限按上限缴纳"],
        [4, "公积金比例合规", "比例在 5%-12% 之间", housing_rate, 0.05, f'=IF(AND(D5>=E5,D5<=0.12),"PASS","WARN")', "比例需符合当地政策"],
        [5, "绩效工资占比", "绩效工资 ≤ 固定薪酬*50%", perf_std, gross_fixed * 0.5, f'=IF(D6<=E6,"PASS","WARN")', "占比过高影响收入稳定性"],
        [6, "个税合规", "应纳税所得额 ≥ 0", f"='计算示例'!{calc_info['taxable_income_cell']}", 0, f'=IF(D7>=E7,"PASS","WARN")', "应纳税所得额为负需检查"],
        [7, "实发工资为正", "实发工资 > 0", f"='计算示例'!{calc_info['net_salary_cell']}", 0, f'=IF(D8>0,"PASS","BLOCK")', "实发工资必须为正"],
    ]

    for i, row in enumerate(rows, 2):
        for j, v in enumerate(row, 1):
            cell = ws.cell(row=i, column=j, value=v)
            is_formula = isinstance(v, str) and v.startswith("=")
            cell.font = styles["green_font"] if is_formula else styles["black_font"]
            cell.border = styles["thin_border"]
            cell.alignment = styles["center_align"]
            if j == 5 and i == 5:
                cell.number_format = "0%"

    calc_ws = wb["计算示例"]
    last_row = 40
    calc_ws.cell(row=last_row, column=1, value="合规状态")
    calc_ws.cell(row=last_row, column=2, value='=COUNTIF(合规校验!F:F,"PASS")&"项通过/"&COUNTA(合规校验!F:F)-1&"项校验"')
    calc_ws.cell(row=last_row, column=2).font = styles["green_font"]

    for col in range(1, len(headers) + 1):
        ws.column_dimensions[get_column_letter(col)].width = 16
    ws.column_dimensions["B"].width = 22
    ws.column_dimensions["C"].width = 24
    ws.column_dimensions["G"].width = 26


# ---------------------------------------------------------------------------
# Word / PDF output
# ---------------------------------------------------------------------------
def _compute_summary_values(params):
    base = params.get("base_salary", 0)
    pos = params.get("position_salary", 0)
    perf_std = params.get("performance_std", 0)
    full = params.get("full_attendance", 0)
    meal = params.get("meal_subsidy", 0)
    trans = params.get("transport_subsidy", 0)
    comm = params.get("communication_subsidy", 0)
    overtime = params.get("overtime_pay", 0)
    actual = params.get("actual_attendance_days", 21)
    scheduled = params.get("scheduled_days", 21.75)
    social_base = params.get("social_base", 0)
    housing_rate = params.get("housing_fund_rate", 0.05)

    kpi_weights = params.get("kpi_weights", [["", 0, 0]])
    avg_score = sum(w[2] for w in kpi_weights) / len(kpi_weights)
    perf_coeff = max(0.6, min(1.2, avg_score / 100))
    perf_actual = round(perf_std * perf_coeff, 2)
    meal_actual = round(meal / scheduled * actual, 2)

    gross = base + pos + perf_actual + overtime + full + meal_actual + trans + comm
    social = social_base * (0.08 + 0.02 + 0.005) + social_base * housing_rate
    taxable = max(gross - social - 5000, 0)
    if taxable <= 0:
        tax = 0
    elif taxable <= 3000:
        tax = taxable * 0.03
    elif taxable <= 12000:
        tax = taxable * 0.10 - 210
    elif taxable <= 25000:
        tax = taxable * 0.20 - 1410
    else:
        tax = taxable * 0.25 - 2660
    tax = round(tax, 2)
    net = round(gross - social - tax, 2)

    sales_commission = params.get("bonus", {}).get("sales_commission", 0)
    annual_fixed = round((base + pos) * 12, 2)
    annual_variable = round((sales_commission + perf_actual) * 12, 2)
    annual_total = round(annual_fixed + annual_variable + (base + pos) + (meal_actual + trans + comm + full) * 12 + 1200, 2)

    return {
        "base": base, "position": pos, "performance_actual": perf_actual,
        "full_attendance": full, "meal_actual": meal_actual, "transport": trans,
        "communication": comm, "overtime": overtime, "gross": gross,
        "social": social, "taxable": taxable, "tax": tax, "net": net,
        "annual_fixed": annual_fixed, "annual_variable": annual_variable, "annual_total": annual_total,
    }


def generate_word(params, path):
    if not HAVE_DOCX:
        return
    doc = Document()
    doc.add_heading(f"{params.get('industry', '通用行业')} · {params.get('position', '岗位')} 薪资模板", 0)
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("【人工智能生成-需人工核验】本模板仅供参考，具体薪酬决策请以当地法律法规及公司政策为准。")

    vals = _compute_summary_values(params)

    def add_table(title, headers, rows):
        doc.add_heading(title, level=1)
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = str(h)
        for row in rows:
            cells = table.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = str(v)
        doc.add_paragraph()

    base = params.get("base_salary", 0)
    pos = params.get("position_salary", 0)
    perf = params.get("performance_std", 0)
    full = params.get("full_attendance", 0)
    meal = params.get("meal_subsidy", 0)
    trans = params.get("transport_subsidy", 0)
    comm = params.get("communication_subsidy", 0)
    overtime = params.get("overtime_pay", 0)

    add_table("一、薪资结构", ["序号", "薪资项目", "项目类型", "金额(元)"], [
        [1, "基本工资", "固定薪资", base],
        [2, "岗位工资", "固定薪资", pos],
        [3, "绩效工资", "浮动薪资", perf],
        [4, "加班工资", "浮动薪资", overtime],
        [5, "全勤奖", "津贴补贴", full],
        [6, "餐补", "津贴补贴", vals["meal_actual"]],
        [7, "交通补贴", "津贴补贴", trans],
        [8, "通讯补贴", "津贴补贴", comm],
        [9, "五险一金(个人)", "代扣项", round(vals["social"], 2)],
        [10, "个人所得税", "代扣项", vals["tax"]],
        [11, "应发工资", "汇总", round(vals["gross"], 2)],
        [12, "实发工资", "汇总", vals["net"]],
    ])

    kpi_rows = [[i + 1, name, f"{weight:.0%}", score] for i, (name, weight, score) in enumerate(params.get("kpi_weights", []))]
    add_table("二、绩效考核", ["序号", "考核维度", "权重", "得分"], kpi_rows)

    add_table("三、调薪机制", ["序号", "调薪类型", "适用条件", "调薪幅度"], [
        [1, "试用期转正调薪", "试用期考核通过", "0-10%"],
        [2, "年度普调", "入职满1年且年度绩效≥B", "3-8%"],
        [3, "晋升调薪", "职级晋升", "10-25%"],
        [4, "绩效调薪", "连续2个季度绩效A或年度绩效A", "5-15%"],
        [5, "市场薪酬调整", "岗位市场薪酬显著变化", "5-10%"],
        [6, "特殊贡献调薪", "重大项目突破/特殊贡献", "10-30%"],
        [7, "留任调薪", "核心人才面临外部高薪挖角", "10-20%"],
    ])

    bb = params.get("bonus_bandwidth", {})
    q_min, q_mid, q_max = bb.get("quarter", (0.5, 0.8, 1.5))
    a_min, a_mid, a_max = bb.get("annual", (1.0, 1.5, 3.0))
    bonus = params.get("bonus", {})
    add_table("四、奖金方案", ["序号", "奖金类型", "计算基数", "计提比例/金额"], [
        [1, "月度销售提成", "个人负责销售额", bonus.get("sales_commission", 0)],
        [2, "月度绩效奖金", "绩效标准金额", f"0-120%"],
        [3, "季度奖金", "月基本工资", f"{q_min}-{q_max}个月"],
        [4, "年度奖金", "月基本工资", f"{a_min}-{a_max}个月"],
        [5, "项目奖金", "项目利润", "5%-15%"],
        [6, "全勤奖", "固定金额", full],
        [7, "伯乐奖", "固定金额", "500-2000元"],
        [8, "年终分红", "公司年度利润", "0.5%-2%"],
    ])

    add_table("五、福利明细", ["序号", "福利项目", "福利类型", "标准/金额"], [
        [1, "养老保险", "法定福利", "基数*16%(单位)+8%(个人)"],
        [2, "医疗保险", "法定福利", "基数*8%(单位)+2%(个人)"],
        [3, "失业保险", "法定福利", "基数*0.5%(单位)+0.5%(个人)"],
        [4, "工伤保险", "法定福利", "基数*0.2%-1.4%(单位)"],
        [5, "住房公积金", "法定福利", "基数*5%-12%(双边)"],
        [6, "补充医疗保险", "补充福利", "约200元/人/年"],
        [7, "带薪年假", "法定福利", "5-15天/年"],
        [8, "节日福利", "补充福利", "300-500元/节"],
        [9, "交通补助", "津贴补贴", f"{trans}元/月"],
        [10, "餐补", "津贴补贴", f"{meal}元/月"],
        [11, "通讯补贴", "津贴补贴", f"{comm}元/月"],
    ])

    add_table("六、计算示例", ["项目", "金额(元/月)"], [
        ["基本工资", base],
        ["岗位工资", pos],
        ["绩效工资", vals["performance_actual"]],
        ["应发工资", round(vals["gross"], 2)],
        ["五险一金(个人)", round(vals["social"], 2)],
        ["个人所得税", vals["tax"]],
        ["实发工资", vals["net"]],
        ["全年总现金收入", vals["annual_total"]],
    ])

    doc.add_paragraph("【人工智能生成-需人工核验】")
    doc.save(path)


def generate_pdf(params, path):
    if not HAVE_REPORTLAB:
        return
    vals = _compute_summary_values(params)
    doc = SimpleDocTemplate(path, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph(f"{params.get('industry', '通用行业')} · {params.get('position', '岗位')} 薪资模板", styles["Title"]))
    story.append(Paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}", styles["Normal"]))
    story.append(Spacer(1, 12))

    data = [
        ["项目", "金额(元/月)"],
        ["基本工资", params.get("base_salary", 0)],
        ["岗位工资", params.get("position_salary", 0)],
        ["绩效工资", vals["performance_actual"]],
        ["应发工资", round(vals["gross"], 2)],
        ["五险一金(个人)", round(vals["social"], 2)],
        ["个人所得税", vals["tax"]],
        ["实发工资", vals["net"]],
        ["全年总现金收入", vals["annual_total"]],
    ]
    t = Table(data, colWidths=[240, 160])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E78")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
    ]))
    story.append(t)
    story.append(Spacer(1, 12))
    story.append(Paragraph("【人工智能生成-需人工核验】本模板仅供参考，具体薪酬决策请以当地法律法规及公司政策为准。", styles["Normal"]))
    doc.build(story)


# ---------------------------------------------------------------------------
# History / compare
# ---------------------------------------------------------------------------
def _params_summary(params):
    return {
        "industry": params.get("industry", ""),
        "position": params.get("position", ""),
        "city": params.get("city", ""),
        "level": params.get("job_level", ""),
        "base_salary": params.get("base_salary", 0),
        "position_salary": params.get("position_salary", 0),
        "performance_std": params.get("performance_std", 0),
        "full_attendance": params.get("full_attendance", 0),
        "meal_subsidy": params.get("meal_subsidy", 0),
        "transport_subsidy": params.get("transport_subsidy", 0),
        "communication_subsidy": params.get("communication_subsidy", 0),
        "social_base": params.get("social_base", 0),
        "housing_fund_rate": params.get("housing_fund_rate", 0.05),
        "sales_commission": params.get("bonus", {}).get("sales_commission", 0),
        "quarter_bonus_months": params.get("bonus", {}).get("quarter_bonus_months", 0.5),
    }


def save_to_history(output_path, params):
    HISTORY_DIR.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_name = Path(output_path).name
    hist_name = f"{ts}_{out_name}"
    hist_path = HISTORY_DIR / hist_name
    shutil.copy2(output_path, hist_path)

    manifest = {
        "timestamp": ts,
        "generated_file": hist_name,
        "position": params.get("position", ""),
        "industry": params.get("industry", ""),
        "city": params.get("city", ""),
        "level": params.get("job_level", ""),
        "params": _params_summary(params),
    }
    manifest_path = HISTORY_DIR / f"{ts}_{Path(out_name).stem}.json"
    with open(manifest_path, "w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)
    return hist_path, manifest_path


def compare_history(position=None):
    HISTORY_DIR.mkdir(parents=True, exist_ok=True)
    manifests = []
    for p in HISTORY_DIR.glob("*.json"):
        try:
            data = json.loads(p.read_text(encoding="utf-8"))
            manifests.append(data)
        except Exception:
            continue
    if not manifests:
        print("未找到历史记录。")
        return

    manifests.sort(key=lambda x: x.get("timestamp", ""), reverse=True)
    if position is None:
        position = manifests[0].get("position")
    filtered = [m for m in manifests if m.get("position") == position]
    if len(filtered) < 2:
        print(f"岗位 '{position}' 的历史记录不足 2 条，无法对比。")
        return

    a, b = filtered[0], filtered[1]
    pa, pb = a.get("params", {}), b.get("params", {})
    keys = ["base_salary", "position_salary", "performance_std", "full_attendance",
            "meal_subsidy", "transport_subsidy", "communication_subsidy", "social_base"]
    print(f"历史版本对比 — 岗位：{position}")
    print(f"  新版本：{a.get('timestamp')}  旧版本：{b.get('timestamp')}")
    print("-" * 50)
    for k in keys:
        va, vb = pa.get(k, 0), pb.get(k, 0)
        diff = va - vb
        print(f"  {k:20s}: {va:>10} -> {vb:>10} (变化 {diff:+.2f})")
    print("-" * 50)


# ---------------------------------------------------------------------------
# Output path helpers
# ---------------------------------------------------------------------------
def generated_filename(params, ext=".xlsx"):
    date = datetime.now().strftime("%Y%m%d")
    base = _safe_filename(f"prajna_薪资模板_{params.get('industry', '通用')}_{params.get('position', '岗位')}_{date}")
    return f"{base}{ext}"


def resolve_output_path(args, params):
    fname = generated_filename(params, ".xlsx")
    if args.output:
        p = Path(args.output).expanduser()
        if p.is_dir():
            return p / fname
        if p.suffix.lower() != ".xlsx":
            p = p.with_suffix(".xlsx")
        return p
    return HISTORY_DIR / fname


# ---------------------------------------------------------------------------
# Batch mode
# ---------------------------------------------------------------------------
def run_batch(args):
    city_data = load_json(DATA_DIR / "city_salary_multipliers.json")
    level_data = load_json(DATA_DIR / "job_level_bandwidth.json")
    kpi_data = load_json(DATA_DIR / "position_kpi_templates.json")
    salary_data = load_json(DATA_DIR / "position_salary_templates.json")

    output_dir = Path(args.output).expanduser() if args.output else HISTORY_DIR
    output_dir.mkdir(parents=True, exist_ok=True)

    summary_rows = []
    generated_files = []
    with open(args.batch, newline="", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        for idx, row in enumerate(reader, 1):
            fields = {
                "industry": (row.get("industry") or "").strip(),
                "position": (row.get("position") or "").strip(),
                "city": (row.get("city") or "").strip(),
                "level": (row.get("level") or "").strip(),
            }
            params = build_params_from_fields(fields, city_data, level_data, kpi_data, salary_data)
            out_name = (row.get("output_name") or "").strip()
            if not out_name:
                out_name = generated_filename(params, ".xlsx")
            elif not out_name.lower().endswith(".xlsx"):
                out_name += ".xlsx"
            out_path = output_dir / out_name

            gen = build_salary_workbook(params, str(out_path), args.format)
            hist_path, _ = save_to_history(gen["excel"], params)
            generated_files.append(hist_path)

            vals = _compute_summary_values(params)
            summary_rows.append([
                idx,
                params.get("industry", ""),
                params.get("position", ""),
                params.get("city", ""),
                params.get("job_level", ""),
                hist_path.name,
                vals["annual_fixed"],
                vals["annual_variable"],
                vals["annual_total"],
            ])
            print(f"  [{idx}] {params.get('industry')} · {params.get('position')} -> {hist_path}")

    if not summary_rows:
        print("CSV 为空，未生成任何文件。")
        return

    summary_path = output_dir / f"prajna_薪资模板_批量汇总_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "批量汇总"
    headers = ["序号", "行业", "岗位", "城市", "职级", "历史文件名", "年固定薪资", "年浮动奖金", "全年总现金收入"]
    styles = _base_styles()
    _set_header(ws, 1, headers, styles)
    for r, row in enumerate(summary_rows, 2):
        _write_row(ws, r, row, styles)
    for col in range(1, len(headers) + 1):
        ws.column_dimensions[get_column_letter(col)].width = 18
    wb.save(summary_path)
    print(f"\n批量汇总已保存：{summary_path}")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(
        description="prajna-salary-template v2.0.0 — 生成完整薪资结构模板",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("--preset", help="使用内置预设")
    parser.add_argument("--industry", help="行业")
    parser.add_argument("--position", help="岗位")
    parser.add_argument("--city", help="城市（应用城市薪酬系数）")
    parser.add_argument("--level", help="职级（P1-P9，应用职级带宽）")
    parser.add_argument("--description", help="自然语言岗位描述")
    parser.add_argument("--input", help="现有薪酬表 xlsx 路径（自动补全）")
    parser.add_argument("--output", help="输出路径或目录（默认 ~/.prajna/salary_template_history/）")
    parser.add_argument("--format", choices=["excel", "word", "pdf", "all"], default="excel",
                        help="输出格式（excel/word/pdf/all）")
    parser.add_argument("--batch", help="批量 CSV 路径（列：industry,position,city,level,output_name）")
    parser.add_argument("--history", action="store_true", help="保存到历史版本库")
    parser.add_argument("--compare", action="store_true", help="对比最近两份同岗位历史记录")
    parser.add_argument("--compare-position", help="指定对比的岗位名称（配合 --compare）")
    parser.add_argument("--list-presets", action="store_true", help="列出所有预设")
    args = parser.parse_args()

    if args.list_presets:
        print("内置预设：")
        for name in list_presets():
            print(f"  - {name}")
        return 0

    if args.compare:
        compare_history(args.compare_position)
        return 0

    if args.batch:
        run_batch(args)
        return 0

    params, recognized, diagnostics = build_params(args)

    if recognized:
        print("从输入文件识别到以下字段：")
        for field, raw in diagnostics.get("recognized", []):
            print(f"  - {field}: {raw}")
        if diagnostics.get("missing"):
            print("未识别字段：", ", ".join(diagnostics["missing"]))

    output_path = resolve_output_path(args, params)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    generated = build_salary_workbook(params, str(output_path), args.format)
    print(f"已生成：{generated['excel']}")
    if "word" in generated:
        print(f"已生成 Word：{generated['word']}")
    if "pdf" in generated:
        print(f"已生成 PDF：{generated['pdf']}")

    if args.history:
        hist_path, manifest_path = save_to_history(generated["excel"], params)
        print(f"已保存历史版本：{hist_path}")
        print(f"  清单：{manifest_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
